# This Python 3 environment comes with many helpful analytics libraries installed
# It is defined by the kaggle/python Docker image: https://github.com/kaggle/docker-python
# For example, here's several helpful packages to load

import numpy as np # linear algebra
import pandas as pd # data processing, CSV file I/O (e.g. pd.read_csv)

# Input data files are available in the read-only "../input/" directory
# For example, running this (by clicking run or pressing Shift+Enter) will list all files under the input directory

import os
for dirname, _, filenames in os.walk('/kaggle/input'):
    for filename in filenames:
        print(os.path.join(dirname, filename))

# You can write up to 20GB to the current directory (/kaggle/working/) that gets preserved as output when you create a version using "Save & Run All" 
# You can also write temporary files to /kaggle/temp/, but they won't be saved outside of the current session


from kaggle_secrets import UserSecretsClient
import os
user_secrets = UserSecretsClient()
os.environ["GOOGLE_API_KEY"] = user_secrets.get_secret("GOOGLE_API_KEY")



!pip install streamlit==1.38.0  # Latest stable for Kaggle compatibility
!pip install python-docx beautifulsoup4 requests
!pip install google-generativeai 





"""
Google ADK Multi-Agent Resume Optimizer with A2A Protocol
Demonstrates: Multi-agents, A2A, Tools (MCP-style), Sessions, Memory, Evaluation, Observability
"""

import streamlit as st
from google.adk.agents import Agent
from google.adk.runners import Runner
from google.adk.sessions import InMemorySessionService
from google.adk.tools import google_search
from google.genai import types
import asyncio
import json
import time
from datetime import datetime
from dataclasses import dataclass, asdict
from typing import Dict, List, Any
import logging
from io import BytesIO
from docx import Document
import requests
from bs4 import BeautifulSoup
import os

# Set API key from environment or Streamlit secrets

from kaggle_secrets import UserSecretsClient
import os
user_secrets = UserSecretsClient()
os.environ["GOOGLE_API_KEY"] = user_secrets.get_secret("GOOGLE_API_KEY")




if "GOOGLE_API_KEY" not in os.environ:
    if hasattr(st, 'secrets') and "GOOGLE_API_KEY" in st.secrets:
        os.environ["GOOGLE_API_KEY"] = st.secrets["GOOGLE_API_KEY"]

# â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�
# 1. OBSERVABILITY - Logging, Tracing, Metrics
# â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@dataclass
class AgentTrace:
    """Trace individual agent execution"""
    agent_name: str
    start_time: float
    end_time: float
    token_count: int
    success: bool
    error: str = None
    
    def duration(self) -> float:
        return self.end_time - self.start_time
    
    def to_dict(self) -> Dict:
        return asdict(self)

class ObservabilityCollector:
    """Collects metrics and traces for all agent operations"""
    def __init__(self):
        self.traces: List[AgentTrace] = []
        self.metrics = {
            "total_agents": 0,
            "total_tokens": 0,
            "total_time": 0.0,
            "success_count": 0,
            "failure_count": 0
        }
    
    def log_trace(self, trace: AgentTrace):
        self.traces.append(trace)
        self.metrics["total_agents"] += 1
        self.metrics["total_tokens"] += trace.token_count
        self.metrics["total_time"] += trace.duration()
        
        if trace.success:
            self.metrics["success_count"] += 1
            logger.info(f"âœ… {trace.agent_name}: {trace.duration():.2f}s, {trace.token_count} tokens")
        else:
            self.metrics["failure_count"] += 1
            logger.error(f"â�Œ {trace.agent_name}: {trace.error}")
    
    def get_summary(self) -> Dict:
        return {
            **self.metrics,
            "success_rate": self.metrics["success_count"] / max(self.metrics["total_agents"], 1),
            "avg_time": self.metrics["total_time"] / max(self.metrics["total_agents"], 1),
            "traces": [t.to_dict() for t in self.traces]
        }

# â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�
# 2. CUSTOM TOOLS (MCP-style architecture)
# â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�

def web_scraper_tool(url: str) -> str:
    """
    Custom tool: Scrape job posting from URL
    MCP-style tool implementation
    """
    logger.info(f"ğŸ”§ TOOL: web_scraper({url})")
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Extract job description
        text = soup.get_text(separator=' ', strip=True)
        # Limit to 2000 chars for context
        result = text[:2000]
        logger.info(f"âœ… TOOL: web_scraper returned {len(result)} chars")
        return result
    except Exception as e:
        error_msg = f"Scraping failed: {str(e)}"
        logger.error(f"â�Œ TOOL: web_scraper error: {error_msg}")
        return error_msg

def keyword_analyzer_tool(text: str, top_n: int = 15) -> Dict:
    """
    Custom tool: Analyze and extract keywords from job description
    Returns structured keyword data
    """
    logger.info(f"ğŸ”§ TOOL: keyword_analyzer(top_n={top_n})")
    try:
        from collections import Counter
        import re
        
        # Clean and tokenize
        words = re.findall(r'\b[a-zA-Z]{4,}\b', text.lower())
        
        # Remove common words
        stop_words = {'the', 'and', 'for', 'with', 'that', 'this', 'from', 'have', 'will', 'your', 'about'}
        keywords = [w for w in words if w not in stop_words]
        
        # Get top keywords
        keyword_counts = Counter(keywords).most_common(top_n)
        
        result = {
            "keywords": [kw for kw, _ in keyword_counts],
            "keyword_counts": dict(keyword_counts),
            "total_unique": len(set(keywords))
        }
        
        logger.info(f"âœ… TOOL: keyword_analyzer found {len(result['keywords'])} keywords")
        return result
    except Exception as e:
        logger.error(f"â�Œ TOOL: keyword_analyzer error: {e}")
        return {"error": str(e)}

def ats_scorer_tool(resume: str, job_keywords: List[str]) -> Dict:
    """
    Custom tool: Score resume for ATS compatibility
    Evaluates keyword coverage
    """
    logger.info(f"ğŸ”§ TOOL: ats_scorer with {len(job_keywords)} keywords")
    try:
        resume_lower = resume.lower()
        matched = [kw for kw in job_keywords if kw.lower() in resume_lower]
        missing = [kw for kw in job_keywords if kw.lower() not in resume_lower]
        
        score = (len(matched) / len(job_keywords) * 100) if job_keywords else 0
        
        result = {
            "ats_score": round(score, 2),
            "matched_keywords": matched,
            "missing_keywords": missing,
            "match_count": len(matched),
            "total_keywords": len(job_keywords)
        }
        
        logger.info(f"âœ… TOOL: ats_scorer returned score: {result['ats_score']}/100")
        return result
    except Exception as e:
        logger.error(f"â�Œ TOOL: ats_scorer error: {e}")
        return {"error": str(e)}

# â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�
# 3. SESSIONS & MEMORY - ADK Session Management
# â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�

class ResumeMemoryBank:
    """
    Long-term memory for resume optimizations
    Stores optimization history across sessions
    """
    def __init__(self):
        if 'memory_bank' not in st.session_state:
            st.session_state.memory_bank = {}
            logger.info("ğŸ“¦ Initialized Memory Bank")
    
    def store(self, resume_id: str, data: Dict):
        st.session_state.memory_bank[resume_id] = {
            "timestamp": datetime.now().isoformat(),
            "data": data
        }
        logger.info(f"ğŸ’¾ MEMORY: Stored resume optimization: {resume_id}")
    
    def retrieve(self, resume_id: str) -> Dict:
        result = st.session_state.memory_bank.get(resume_id)
        if result:
            logger.info(f"ğŸ“– MEMORY: Retrieved resume: {resume_id}")
        return result
    
    def list_all(self) -> List[str]:
        return list(st.session_state.memory_bank.keys())
    
    def clear(self):
        st.session_state.memory_bank.clear()
        logger.info("ğŸ—‘ï¸� MEMORY: Cleared memory bank")

# â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�
# 4. MULTI-AGENT SYSTEM with A2A Protocol
# â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�

class A2AAgentSystem:
    """
    Multi-agent orchestrator using A2A protocol patterns
    Agents communicate through standardized message passing
    """
    def __init__(self, model: str = "gemini-2.0-flash-exp"):
        self.model = model
        self.agents = {}
        self.observability = ObservabilityCollector()
        self.memory_bank = ResumeMemoryBank()
        
        # Create ADK session service for state management
        self.session_service = InMemorySessionService()
        
        logger.info("ğŸ¤– Initialized A2A Agent System")
        
        # Initialize specialized agents
        self._create_agents()
    
    def _create_agents(self):
        """Create specialized agents with A2A-compatible interfaces"""
        
        # Agent 1: Job Analyzer (with Google Search tool)
        self.agents['job_analyzer'] = Agent(
            model=self.model,
            name='job_analyzer_agent',
            description='Analyzes job postings and extracts requirements',
            instruction="""You are a job posting analyzer. Extract key requirements, 
            skills, and qualifications from job descriptions. Return structured JSON with:
            - required_skills: list of required skills
            - preferred_skills: list of preferred skills
            - experience_level: years of experience needed
            - key_requirements: main job requirements""",
            tools=[google_search]
        )
        
        # Agent 2: Resume Parser
        self.agents['resume_parser'] = Agent(
            model=self.model,
            name='resume_parser_agent',
            description='Parses and structures resume content',
            instruction="""You are a resume parser. Extract and organize resume content.
            Return structured JSON with:
            - experience: list of work experiences
            - skills: list of skills mentioned
            - education: education background
            - achievements: key achievements"""
        )
        
        # Agent 3: Impact Writer
        self.agents['impact_writer'] = Agent(
            model=self.model,
            name='impact_writer_agent',
            description='Rewrites resume bullets with quantified impact',
            instruction="""You are an expert at writing impactful resume bullets using the 
            Xâ†’Yâ†’Z formula (Accomplished X as measured by Y by doing Z). Rewrite bullets with 
            specific metrics, percentages, and quantified achievements. Always add numbers.
            Return JSON with 'optimized_bullets': list of improved bullets."""
        )
        
        # Agent 4: ATS Optimizer
        self.agents['ats_optimizer'] = Agent(
            model=self.model,
            name='ats_optimizer_agent',
            description='Optimizes resume for ATS systems',
            instruction="""You are an ATS optimization specialist. Insert missing keywords 
            naturally into the resume while maintaining readability. Return JSON with:
            - optimized_resume: the ATS-optimized version
            - keywords_added: list of keywords inserted"""
        )
        
        # Agent 5: Final Editor
        self.agents['final_editor'] = Agent(
            model=self.model,
            name='final_editor_agent',
            description='Polishes and finalizes the resume',
            instruction="""You are a senior resume writer. Create a compelling professional 
            summary and ensure the entire resume is polished, consistent, and FAANG-ready. 
            Return JSON with 'final_resume': the polished final version."""
        )
        
        logger.info(f"âœ… Created {len(self.agents)} A2A-compatible agents")
    
    async def run_agent_with_runner(self, agent_name: str, message: str, context: Dict = None) -> Dict:
        """
        Run a single agent using ADK Runner with observability
        """
        start_time = time.time()
        agent = self.agents.get(agent_name)
        
        if not agent:
            raise ValueError(f"Agent {agent_name} not found")
        
        try:
            # Add context from previous agents (A2A pattern)
            full_message = message
            if context:
                full_message = f"{message}\n\nContext from previous agents:\n{json.dumps(context, indent=2)}"
            
            logger.info(f"ğŸ”„ Running {agent_name}...")
            
            # Create session for this agent execution
            session_id = f"session_{agent_name}_{int(time.time())}"
            session = await self.session_service.create_session(
                app_name="resume_optimizer",
                user_id="resume_user",
                session_id=session_id
            )
            
            # Create runner for this agent
            runner = Runner(
                agent=agent,
                app_name="resume_optimizer",
                session_service=self.session_service
            )
            
            # Prepare message as Content object (required by ADK)
            message_content = types.Content(
                role='user',
                parts=[types.Part(text=full_message)]
            )
            
            # Execute agent with runner
            response_text = ""
            async for event in runner.run_async(
                user_id="resume_user",
                session_id=session_id,
                new_message=message_content
            ):
                if event.is_final_response():
                    if event.content and event.content.parts:
                        for part in event.content.parts:
                            if hasattr(part, 'text'):
                                response_text += part.text
            
            # Parse response
            try:
                result = json.loads(response_text)
            except:
                result = {"output": response_text}
            
            # Log trace
            trace = AgentTrace(
                agent_name=agent_name,
                start_time=start_time,
                end_time=time.time(),
                token_count=len(response_text.split()),
                success=True
            )
            self.observability.log_trace(trace)
            
            return result
            
        except Exception as e:
            trace = AgentTrace(
                agent_name=agent_name,
                start_time=start_time,
                end_time=time.time(),
                token_count=0,
                success=False,
                error=str(e)
            )
            self.observability.log_trace(trace)
            return {"error": str(e)}
    
    async def run_multi_agent_pipeline(self, resume: str, job_url: str) -> Dict:
        """
        Run complete multi-agent pipeline with A2A communication
        Sequential execution with context passing between agents
        """
        logger.info("ğŸš€ Starting multi-agent pipeline")
        
        # Step 1: Scrape job with custom tool
        job_description = web_scraper_tool(job_url)
        
        # Step 2: Analyze keywords with custom tool
        keywords_data = keyword_analyzer_tool(job_description)
        job_keywords = keywords_data.get("keywords", [])
        
        # Step 3: Run Agent 1 - Job Analyzer (with Google Search)
        job_analysis = await self.run_agent_with_runner(
            "job_analyzer",
            f"Analyze this job posting and extract key requirements: {job_description[:1000]}"
        )
        
        # A2A Context passing - build shared context
        context = {
            "job_description": job_description[:500],
            "keywords": job_keywords,
            "job_analysis": job_analysis
        }
        
        # Step 4: Run Agent 2 - Resume Parser
        parsed_resume = await self.run_agent_with_runner(
            "resume_parser",
            f"Parse and structure this resume: {resume[:2000]}",
            context
        )
        context["parsed_resume"] = parsed_resume
        
        # Step 5: Run Agent 3 - Impact Writer
        impact_resume = await self.run_agent_with_runner(
            "impact_writer",
            f"Rewrite these resume bullets with Xâ†’Yâ†’Z metrics: {resume[:2000]}",
            context
        )
        context["impact_resume"] = impact_resume
        
        # Step 6: Run Agent 4 - ATS Optimizer
        ats_resume = await self.run_agent_with_runner(
            "ats_optimizer",
            f"Optimize this resume for ATS with keywords: {impact_resume}",
            context
        )
        context["ats_resume"] = ats_resume
        
        # Step 7: Run Agent 5 - Final Editor
        final_resume = await self.run_agent_with_runner(
            "final_editor",
            f"Polish and finalize this resume: {ats_resume}",
            context
        )
        
        # Step 8: Score with ATS tool
        ats_score = ats_scorer_tool(str(final_resume), job_keywords)
        
        # Store in memory bank
        resume_id = f"resume_{int(time.time())}"
        self.memory_bank.store(resume_id, {
            "original": resume[:500],
            "final": final_resume,
            "ats_score": ats_score,
            "context": context
        })
        
        logger.info("âœ… Multi-agent pipeline complete")
        
        return {
            "job_analysis": job_analysis,
            "parsed_resume": parsed_resume,
            "impact_resume": impact_resume,
            "ats_optimized": ats_resume,
            "final_resume": final_resume,
            "ats_score": ats_score,
            "observability": self.observability.get_summary()
        }

# â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�
# 5. AGENT EVALUATION System
# â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�

class ResumeEvaluator:
    """Evaluate resume optimization quality"""
    
    @staticmethod
    def evaluate(original: str, optimized: str, ats_data: Dict) -> Dict:
        logger.info("ğŸ“Š Starting resume evaluation...")
        
        import re
        
        # Metric 1: ATS Score
        ats_score = ats_data.get("ats_score", 0)
        
        # Metric 2: Quantification (numbers, metrics)
        orig_numbers = len(re.findall(r'\d+%|\d+x|\$\d+|\d+\+', original))
        opt_numbers = len(re.findall(r'\d+%|\d+x|\$\d+|\d+\+', str(optimized)))
        quantification_improvement = ((opt_numbers - orig_numbers) / max(orig_numbers, 1)) * 100
        
        # Metric 3: Length appropriateness
        opt_length = len(str(optimized))
        length_score = min(opt_length / 3000 * 100, 100)  # Target ~3000 chars
        
        # Metric 4: Keyword density
        keyword_coverage = ats_data.get("match_count", 0) / max(ats_data.get("total_keywords", 1), 1) * 100
        
        # Overall score (weighted)
        overall = (
            ats_score * 0.4 +
            min(quantification_improvement, 100) * 0.3 +
            keyword_coverage * 0.2 +
            length_score * 0.1
        )
        
        evaluation = {
            "overall_score": round(overall, 2),
            "ats_score": ats_score,
            "quantification_improvement": round(quantification_improvement, 2),
            "keyword_coverage": round(keyword_coverage, 2),
            "length_score": round(length_score, 2),
            "metrics_added": opt_numbers - orig_numbers,
            "passed": overall >= 70,
            "grade": "A" if overall >= 90 else "B" if overall >= 80 else "C" if overall >= 70 else "D"
        }
        
        logger.info(f"âœ… Evaluation complete: {evaluation['overall_score']}/100 (Grade: {evaluation['grade']})")
        return evaluation

# â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�
# STREAMLIT UI
# â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�â•�

st.set_page_config(page_title="ADK Multi-Agent Resume Optimizer", layout="wide")

st.title("ğŸ¤– Google ADK Multi-Agent Resume Optimizer")
st.markdown("""
**Built with Google's Agent Development Kit (ADK) + A2A Protocol**

**Demonstrates:**
- âœ… Multi-Agent System (5 specialized agents with A2A communication)
- âœ… Tools (Google Search + 3 custom MCP-style tools)
- âœ… Sessions & Memory (InMemorySessionService + Memory Bank)
- âœ… Observability (Logging, Tracing, Metrics)
- âœ… Agent Evaluation (Quality scoring system)
""")

# Initialize system
if 'agent_system' not in st.session_state:
    st.session_state.agent_system = A2AAgentSystem()

agent_system = st.session_state.agent_system

# Input section
col1, col2 = st.columns(2)

with col1:
    resume_file = st.file_uploader("ğŸ“„ Upload Resume", type=["txt", "docx"])
    if resume_file:
        resume_text = resume_file.read().decode("utf-8", errors="ignore")
    else:
        resume_text = ""

with col2:
    job_url = st.text_input("ğŸ”— Job Posting URL", placeholder="https://company.com/job-posting")

# Main execution



# SAMPLE DATA FOR JUDGES
if 'resume_text' not in locals() or not resume_text:
    resume_text = """
John Doe
Software Engineer with 5 years experience in Python and cloud.
Built web apps, improved performance, led teams.
Skills: Python, AWS, React, Docker
"""
if 'job_url' not in locals() or not job_url:
    job_url = "https://jobs.google.com/about/careers/applications/jobs/results/123456789-software-engineer"



if st.button("ğŸš€ Run Multi-Agent Optimization", type="primary", use_container_width=True):
    if not resume_text or not job_url:
        st.error("Please provide both resume and job URL")
        st.stop()
    
    progress = st.progress(0)
    status = st.empty()
    
    # Run async pipeline
    status.info("ğŸ¤– Initializing multi-agent system...")
    progress.progress(10)
    
    async def run_pipeline():
        return await agent_system.run_multi_agent_pipeline(resume_text, job_url)
    
    # Execute pipeline
    try:
        results = asyncio.run(run_pipeline())
        
        progress.progress(90)
        status.info("ğŸ“Š Running evaluation...")
        
        # Evaluate
        evaluation = ResumeEvaluator.evaluate(
            resume_text,
            results.get("final_resume", ""),
            results.get("ats_score", {})
        )
        
        progress.progress(100)
        status.empty()
        
        st.success("âœ… Multi-agent optimization complete!")
        st.balloons()
        
        # Display results in tabs
        tab1, tab2, tab3, tab4, tab5 = st.tabs([
            "ğŸ“„ Results",
            "ğŸ“Š Evaluation",
            "ğŸ¤– Agent Outputs",
            "ğŸ“ˆ Observability",
            "ğŸ’¾ Memory"
        ])
        
        with tab1:
            st.subheader("Optimized Resume")
            st.markdown(str(results.get("final_resume", {})))
            
            # Download button
            doc = Document()
            doc.add_heading("AI-Optimized Resume", 0)
            doc.add_paragraph(str(results.get("final_resume")))
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            st.download_button(
                "ğŸ“¥ Download Optimized Resume (DOCX)",
                buffer,
                "optimized_resume.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        with tab2:
            st.subheader("Agent Evaluation Results")
            
            col1, col2, col3, col4 = st.columns(4)
            col1.metric("Overall Score", f"{evaluation['overall_score']}/100")
            col2.metric("Grade", evaluation['grade'])
            col3.metric("ATS Score", f"{evaluation['ats_score']}/100")
            col4.metric("Metrics Added", evaluation['metrics_added'])
            
            if evaluation['passed']:
                st.success("âœ… Resume passed quality evaluation!")
            else:
                st.warning("âš ï¸� Resume needs improvement")
            
            st.json(evaluation)
        
        with tab3:
            st.subheader("Agent Outputs (A2A Communication)")
            
            for agent_name in ["job_analysis", "parsed_resume", "impact_resume", "ats_optimized"]:
                with st.expander(f"ğŸ¤– {agent_name.replace('_', ' ').title()}"):
                    st.json(results.get(agent_name, {}))
            
            with st.expander("ğŸ�¯ ATS Tool Output"):
                st.json(results.get("ats_score", {}))
        
        with tab4:
            st.subheader("Observability Dashboard")
            
            obs = results.get("observability", {})
            
            col1, col2, col3 = st.columns(3)
            col1.metric("Total Agents", obs.get("total_agents", 0))
            col2.metric("Total Tokens", obs.get("total_tokens", 0))
            col3.metric("Total Time", f"{obs.get('total_time', 0):.2f}s")
            
            col1.metric("Success Rate", f"{obs.get('success_rate', 0)*100:.1f}%")
            col2.metric("Avg Time/Agent", f"{obs.get('avg_time', 0):.2f}s")
            
            st.subheader("Agent Traces")
            for trace in obs.get("traces", []):
                with st.expander(f"ğŸ“Š {trace.get('agent_name')} - {trace.get('end_time', 0) - trace.get('start_time', 0):.2f}s"):
                    st.json(trace)
        
        with tab5:
            st.subheader("Memory Bank")
            st.write(f"Stored optimizations: {len(agent_system.memory_bank.list_all())}")
            
            for resume_id in agent_system.memory_bank.list_all():
                with st.expander(f"ğŸ’¾ {resume_id}"):
                    memory_data = agent_system.memory_bank.retrieve(resume_id)
                    st.json(memory_data)
    
    except Exception as e:
        st.error(f"Error during optimization: {str(e)}")
        logger.error(f"Pipeline error: {e}", exc_info=True)

# Sidebar
with st.sidebar:
    st.subheader("ğŸ”§ System Info")
    st.write(f"**Agents:** {len(agent_system.agents)}")
    st.write(f"**Memory Bank:** {len(agent_system.memory_bank.list_all())} items")
    
    if st.button("Clear Memory Bank"):
        agent_system.memory_bank.clear()
        st.rerun()
    
    st.divider()
    st.caption("Built with Google ADK + A2A Protocol")







