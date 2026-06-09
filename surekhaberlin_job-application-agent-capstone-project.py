# This Python 3 environment comes with many helpful analytics libraries installed
# It is defined by the kaggle/python Docker image: https://github.com/kaggle/docker-python
# For example, here's several helpful packages to load

import numpy as np # linear algebra
import pandas as pd # data processing, CSV file I/O (e.g. pd.read_csv)

# Input data files are available in the read-only "../input/" directory
# For example, running this (by clicking run or pressing Shift+Enter) will list all files under the input directory

import os
for dirname, _, filenames in os.walk('/kaggle/input'):
    for filename in filenames:
        print(os.path.join(dirname, filename))

# You can write up to 20GB to the current directory (/kaggle/working/) that gets preserved as output when you create a version using "Save & Run All" 
# You can also write temporary files to /kaggle/temp/, but they won't be saved outside of the current session


import warnings
warnings.filterwarnings("ignore", category=RuntimeWarning)



!pip install openai requests pandas



pip install openai==0.28.0



pip install requests beautifulsoup4



import openai
import requests
import pandas as pd
import json
import os


import os, json, logging, requests
from datetime import datetime
import pandas as pd

# Try loading API keys
OPENAI_AVAILABLE, GOOGLE_API_AVAILABLE = False, False
OPENAI_API_KEY, GOOGLE_API_KEY = None, None

try:
    from kaggle_secrets import UserSecretsClient
    user_secrets = UserSecretsClient()

    OPENAI_API_KEY = user_secrets.get_secret("OPENAI_API_KEY")
    if OPENAI_API_KEY:
        import openai
        openai.api_key = OPENAI_API_KEY
        OPENAI_AVAILABLE = True
        print("âœ… OpenAI API key loaded!")

    GOOGLE_API_KEY = user_secrets.get_secret("GOOGLE_API_KEY")
    if GOOGLE_API_KEY:
        GOOGLE_API_AVAILABLE = True
        print("ğŸ”� Google API key loaded!")

except:
    print("âš ï¸� Kaggle secrets unavailable â€” running API-free mode.")

# Logging setup (ALWAYS WORKS)
LOG_PATH = "/kaggle/working/job_assistant.log"
logging.basicConfig(filename=LOG_PATH, level=logging.INFO,
                    format="%(asctime)s [%(levelname)s] %(message)s")
logging.info("Setup complete.")

print("\nğŸ”§ SETUP COMPLETE")
print(f"OpenAI available: {OPENAI_AVAILABLE}")
print(f"Google API available: {GOOGLE_API_AVAILABLE}")



# Create a resume file inside Kaggle Notebook
with open("resume.txt", "w") as f:
    f.write("""Tom Cruise
Data Analyst
Toronto, ON | tom.cruise@email.com | (555) 123-4567 | LinkedIn: https://linkedin.com/in/tomcruise

Professional Summary
Results-driven Data Analyst with 5 years of experience in analyzing complex datasets, generating actionable insights, and supporting strategic decision-making. Skilled in Python, SQL, Power BI, and Excel with a proven ability to deliver data-driven solutions that improve business performance. Adept at building dashboards, automating reports, and collaborating with cross-functional teams.

Key Skills
- Data Analysis & Visualization: Python (Pandas, NumPy, Matplotlib, Seaborn), SQL, Power BI, Tableau, Excel
- Data Management: ETL processes, data cleaning, database querying
- Reporting & Insights: KPI dashboards, automated reporting, trend analysis
- Tools & Technologies: Git, Jupyter Notebook, Google Sheets, Excel VBA

Professional Experience
Data Analyst | XYZ Analytics Inc., Toronto, ON
Jan 2020 â€“ Present
- Developed interactive dashboards in Power BI to track key business metrics, reducing reporting time by 40%.
- Performed data cleaning, transformation, and analysis using Python and SQL to support marketing and finance teams.
- Designed automated Excel reporting templates, improving team efficiency and accuracy.
- Collaborated with stakeholders to translate business requirements into actionable insights, leading to a 15% increase in revenue.

Junior Data Analyst | ABC Corp., Toronto, ON
Jun 2018 â€“ Dec 2019
- Assisted in data collection, cleaning, and analysis for multiple projects, ensuring high-quality datasets for reporting.
- Generated weekly and monthly reports using SQL and Excel for management.
- Supported senior analysts in dashboard creation and visualization, improving decision-making speed.

Education
BSc in Computer Science â€“ University of Toronto, Toronto, ON
Graduated: 2018

Certifications
- Microsoft Certified: Data Analyst Associate (Power BI)
- Python for Data Science (Coursera)

Projects
- Sales Performance Dashboard: Created a Power BI dashboard integrating multiple datasets to track sales KPIs for management.
- Customer Segmentation Analysis: Performed clustering analysis in Python to identify key customer segments for targeted marketing.
""")

print("resume.txt created!")



# Comprehensive skills list for Data Analyst and Data Scientist
skills_text = """
Python, SQL, Excel, Power BI, Tableau, R, Data Cleaning, Data Wrangling, Data Analysis, Data Visualization, 
Matplotlib, Seaborn, Plotly, Dash, MySQL, PostgreSQL, SQL Server, NoSQL, MongoDB, Google Analytics, 
ETL, APIs, Web Scraping, Statistics, Descriptive Statistics, Inferential Statistics, A/B Testing, 
Regression Analysis, Forecasting, Trend Analysis, Machine Learning, Deep Learning, NLP, Time Series Analysis, 
Predictive Modeling, Big Data, Hadoop, Spark, Cloud (AWS, Azure, GCP), Git, Docker, Data Storytelling, Dashboard Design
"""

# Remove extra spaces and line breaks
skills_text_clean = ", ".join([s.strip() for s in skills_text.split(",")])

# Write to skills.txt
with open("skills.txt", "w", encoding="utf-8") as f:
    f.write(skills_text_clean)

print("skills.txt file created successfully!")



import requests
import pandas as pd

def fetch_remote_jobs():
    url = "https://remoteok.com/api"  # Works NOW
    headers = {"User-Agent": "Mozilla/5.0"}  # Prevent blocking
    
    response = requests.get(url, headers=headers)

    if response.status_code != 200:
        return pd.DataFrame([{"ERROR": "Request failed"}])

    data = response.json()[1:]  # first item is API info

    jobs = []
    for job in data:
        jobs.append({
            "Job Title": job.get("position"),
            "Company": job.get("company"),
            "Location": job.get("location", "Remote"),
            "Tags": ", ".join(job.get("tags", [])),
            "URL": job.get("url"),
        })

    return pd.DataFrame(jobs)

df_jobs = fetch_remote_jobs()
df_jobs.head()



df_filtered = df_jobs[df_jobs["Job Title"].str.contains("data", case=False, na=False)]
df_filtered.head()



skills = ["python", "sql", "power bi", "tableau"]

def find_skills(text):
    return ", ".join([s for s in skills if s.lower() in text.lower()]) or "None"

df_filtered["Matched Skills"] = df_filtered["Tags"].apply(find_skills)
df_filtered.head()



df_filtered.to_excel("remote_jobs.xlsx", index=False)
print("Saved: remote_jobs.xlsx")



# Read resume and skills text files
with open("resume.txt", "r", encoding="utf-8") as f:
    resume_text = f.read()

with open("skills.txt", "r", encoding="utf-8") as f:
    skills = [s.strip() for s in f.read().split(",")]

print("Resume loaded:", resume_text[:200], "...")  # preview first 200 chars
print("Skills loaded:", skills)



def skill_match_score(matched_skills, all_skills):
    matched = [s for s in all_skills if s.lower() in matched_skills.lower()]
    return int(len(matched)/len(all_skills) * 100)



df_filtered.columns



import os
os.environ["OPENAI_API_KEY"] = "OPENAI_API_KEY"  # replace with your actual key



import openai
openai.api_key = os.environ.get("OPENAI_API_KEY")



def skill_match_score(matched_skills, resume_skills):
    matched = [s for s in resume_skills if s.lower() in matched_skills.lower()]
    return int(len(matched)/len(resume_skills) * 100)  # 0-100

df_filtered["score"] = df_filtered["Matched Skills"].apply(lambda x: skill_match_score(x, skills))
df_filtered



# Install Gemini SDK and python-docx
!pip install google-genai python-docx --quiet



import os
os.environ["GOOGLE_API_KEY"] = "GOOGLE_API_KEY"



# Suppose you created it like this earlier:
with open("resume.txt", "w") as f:
    f.write("""Tom Cruise
Data Analyst
Toronto, ON | tom.cruise@email.com | (555) 123-4567 | LinkedIn: https://linkedin.com/in/tomcruise

Professional Summary
Results-driven Data Analyst with 5 years of experience in analyzing complex datasets, generating actionable insights, and supporting strategic decision-making. Skilled in Python, SQL, Power BI, and Excel with a proven ability to deliver data-driven solutions that improve business performance. Adept at building dashboards, automating reports, and collaborating with cross-functional teams.

Key Skills
- Data Analysis & Visualization: Python (Pandas, NumPy, Matplotlib, Seaborn), SQL, Power BI, Tableau, Excel
- Data Management: ETL processes, data cleaning, database querying
- Reporting & Insights: KPI dashboards, automated reporting, trend analysis
- Tools & Technologies: Git, Jupyter Notebook, Google Sheets, Excel VBA

Professional Experience
Data Analyst | XYZ Analytics Inc., Toronto, ON
Jan 2020 â€“ Present
- Developed interactive dashboards in Power BI to track key business metrics, reducing reporting time by 40%.
- Performed data cleaning, transformation, and analysis using Python and SQL to support marketing and finance teams.
- Designed automated Excel reporting templates, improving team efficiency and accuracy.
- Collaborated with stakeholders to translate business requirements into actionable insights, leading to a 15% increase in revenue.

Junior Data Analyst | ABC Corp., Toronto, ON
Jun 2018 â€“ Dec 2019
- Assisted in data collection, cleaning, and analysis for multiple projects, ensuring high-quality datasets for reporting.
- Generated weekly and monthly reports using SQL and Excel for management.
- Supported senior analysts in dashboard creation and visualization, improving decision-making speed.

Education
BSc in Computer Science â€“ University of Toronto, Toronto, ON
Graduated: 2018

Certifications
- Microsoft Certified: Data Analyst Associate (Power BI)
- Python for Data Science (Coursera)

Projects
- Sales Performance Dashboard: Created a Power BI dashboard integrating multiple datasets to track sales KPIs for management.
- Customer Segmentation Analysis: Performed clustering analysis in Python to identify key customer segments for targeted marketing.
""")

print("resume.txt created!")

print("âœ… resume.txt created in working directory")



# Make sure the path matches your working directory
resume_file_path = "resume.txt"

# Read the file contents into a variable
with open(resume_file_path, "r", encoding="utf-8") as f:
    resume_text = f.read()

print("âœ… resume.txt loaded into variable 'resume_text'. Preview:")
print(resume_text[:300])



from google import genai
from docx import Document

# Initialize Gemini client
client = genai.Client()

def generate_resume_gemini(resume_text, model="gemini-2.5-flash"):
    """
    Generate a polished, ATS-friendly resume using Google Gemini.
    Saves both .txt and .docx versions.
    """
    prompt = f"""
Rewrite this resume for a Data Analyst / Data Scientist role.
Make it professional, clear, and ATS-friendly:

{resume_text}
"""
    # Generate content from Gemini
    response = client.models.generate_content(
        model=model,
        contents=prompt
    )
    polished_resume = response.text.strip()

    # Save as .txt
    with open("Resume_Gemini.txt", "w", encoding="utf-8") as f:
        f.write(polished_resume)

    # Save as .docx
    doc = Document()
    for line in polished_resume.splitlines():
        doc.add_paragraph(line)
    doc.save("Resume_Gemini.docx")

    print("âœ… Resume generated and saved as Resume_Gemini.txt and Resume_Gemini.docx")
    return polished_resume

# Generate polished resume
polished_resume = generate_resume_gemini(resume_text)
print("Resume Preview:\n", polished_resume[:500])




def generate_cover_letter_gemini(job_title, company, job_desc, resume_text, model="gemini-2.5-flash"):
    """
    Generate a personalized cover letter using Gemini.
    Saves both .txt and .docx files.
    """
    prompt = f"""
Write a professional cover letter (3-4 paragraphs) for the following job:

Job Title: {job_title}
Company: {company}
Job Description:
{job_desc}

Use this resume:
{resume_text}

Return only the cover letter text.
"""

    # Generate content from Gemini
    response = client.models.generate_content(
        model=model,
        contents=prompt
    )
    cover_letter = response.text.strip()

    # Save to .txt
    txt_filename = f"CoverLetter_{job_title}_{company}.txt".replace(" ", "_")[:100]
    with open(txt_filename, "w", encoding="utf-8") as f:
        f.write(cover_letter)

    # Save to .docx
    doc = Document()
    for line in cover_letter.splitlines():
        doc.add_paragraph(line)
    doc_filename = f"CoverLetter_{job_title}_{company}.docx".replace(" ", "_")[:100]
    doc.save(doc_filename)

    print(f"âœ… Cover letter saved as {txt_filename} and {doc_filename}")
    return cover_letter

# Example usage:
job_title = "Data Analyst"
company = "TechCorp"
job_description = "Looking for expert in Python, SQL, data visualization, Power BI. Must analyze large datasets and create reports."

cover_letter = generate_cover_letter_gemini(job_title, company, job_description, resume_text)
print("Cover Letter Preview:\n", cover_letter[:800])



df_filtered["Matched Skills"] = ""    # Fill later if you want
df_filtered["resume_file"] = "Resume_Gemini.txt"
df_filtered["cover_letter_file"] = ""
df_filtered["status"] = "To Apply"
df_filtered["apply_date"] = ""

df_filtered.head()



df_filtered.to_excel("job_applications_tracker.xlsx", index=False)
print("ğŸ“� Tracker saved: job_applications_tracker.xlsx")



df_filtered.loc[0, "status"] = "Applied"
df_filtered.loc[0, "apply_date"] = "2025-02-01"

df_filtered.to_excel("job_applications_tracker.xlsx", index=False)
print("ğŸ“Œ Status updated and saved!")



!pip install ics --quiet



def create_ics(job_title, company, date="20251201T100000"):
    # Convert to string in case of missing values
    job_title = str(job_title) if pd.notna(job_title) else "Unknown_Job"
    company = str(company) if pd.notna(company) else "Unknown_Company"

    content = f"""BEGIN:VCALENDAR
VERSION:2.0
BEGIN:VEVENT
SUMMARY:Follow-up - {job_title} at {company}
DTSTART:{date}
DURATION:PT15M
END:VEVENT
END:VCALENDAR"""

    # File name safe for saving
    filename = f"{job_title.replace(' ', '_')}_{company.replace(' ', '_')}.ics"
    with open(filename, "w") as f:
        f.write(content)

    return filename



# Make a copy to avoid SettingWithCopyWarning
df_filtered = df_filtered.copy()

# Create reminder files for each job
df_filtered.loc[:, "reminder_file"] = df_filtered.apply(
    lambda row: create_ics(row.get("Job Title"), row.get("Company")), axis=1
)

# Check results
df_filtered.head()



from datetime import datetime

df_filtered.loc[:, "Date Added"] = datetime.now().strftime("%Y-%m-%d")
df_filtered.loc[:, "Status"] = "To Apply"



df_filtered["reminder_file"] = df_filtered.apply(
    lambda row: create_ics(row.get("Job Title"), row.get("Company")), axis=1
)



df_filtered = df_filtered.dropna(subset=["Job Title", "Company"])

df_filtered["reminder_file"] = df_filtered.apply(
    lambda row: create_ics(row["Job Title"], row["Company"]), axis=1
)



import pandas as pd

file_path = "job_applications_tracker.xlsx"
df_filtered.to_excel(file_path, index=False)

print("ğŸ“� Excel tracker created:", file_path)



# Ensure df_filtered exists before running this
import pandas as pd

# Add useful tracking columns automatically
df_filtered["Application_Status"] = ""        # e.g., Pending / Applied / Interview / Rejected
df_filtered["Applied_Date"] = ""              # e.g., 2025-11-28
df_filtered["Resume_File"] = "Resume_Generated.txt"
df_filtered["Cover_Letter_File"] = "Cover_Letter_Generated.txt"
df_filtered["Reminder_File"] = ""             # will be filled later with .ics

# ========== CREATE EXCEL & CSV TRACKER ==========
excel_path = "Job_Applications_Tracker.xlsx"
csv_path = "Job_Applications_Tracker.csv"

df_filtered.to_excel(excel_path, index=False)
df_filtered.to_csv(csv_path, index=False)

print("ğŸ“� Excel Tracker Created:", excel_path)
print("ğŸ“� CSV Tracker Created:", csv_path)



import pandas as pd

def create_ics(job_title, company):
    if isinstance(job_title, float) or isinstance(company, float):
        return ""  # Avoid errors

    content = f"""BEGIN:VCALENDAR
VERSION:2.0
BEGIN:VEVENT
SUMMARY:Follow-up - {job_title} at {company}
DTSTART:20251201T100000
DURATION:PT15M
END:VEVENT
END:VCALENDAR"""

    filename = f"{job_title.replace(' ', '_')}_{company.replace(' ', '_')}.ics"
    with open(filename, "w") as f:
        f.write(content)

    return filename

df_filtered["Reminder_File"] = df_filtered.apply(
    lambda row: create_ics(row["Job Title"], row["Company"]), axis=1
)

# Save again after adding reminders:
df_filtered.to_excel("Job_Applications_Tracker.xlsx", index=False)

print("â�° Reminder files created & tracker updated!")



# Display entire DataFrame in Kaggle
pd.set_option('display.max_rows', None)     # show all rows
pd.set_option('display.max_columns', None)  # show all columns
pd.set_option('display.width', None)        # no wrapping

display(df_filtered)  # or just df_filtered



import os

# List all files
files = os.listdir()
for f in files:
    print(f)



df_excel_preview = pd.read_excel("Job_Applications_Tracker.xlsx")
df_excel_preview.head()   # Show top 5 rows



for index, row in df_filtered.iterrows():
    print("Job Title:", row["Job Title"])
    print("Company:", row["Company"])
    print("Reminder File:", row["Reminder_File"])
    print("Resume:", row["Resume_File"])
    print("Cover Letter:", row["Cover_Letter_File"])
    print("-" * 50)






