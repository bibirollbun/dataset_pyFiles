# Instalar as bibliotecas que o Kaggle não tem por defeito
!pip install python-docx openpyxl

import openpyxl 
import os
from typing import List, Dict, Any
from docx import Document
from docx.shared import Pt
import datetime 

# =========================================================================
# FERRAMENTA 1: Leitura do Excel (Esta está correta)
# =========================================================================

def ler_dados_peso_excel_openpyxl(caminho_ficheiro: str) -> List[Dict[str, Any]]:
    """
    Lê o ficheiro .xlsx diretamente com openpyxl.
    Converte as datas de string (com .strip()) para objetos datetime.
    """
    dados_combinados = []
    
    try:
        workbook = openpyxl.load_workbook(caminho_ficheiro, data_only=True)
        
        for nome_jogador in workbook.sheetnames:
            sheet = workbook[nome_jogador]
            
            for row in sheet.iter_rows(min_row=2):
                data_hora_val = row[0].value
                peso_kg_val = row[1].value
                
                if data_hora_val is None or peso_kg_val is None:
                    continue
                
                data_hora_obj = None
                try:
                    if isinstance(data_hora_val, datetime.datetime):
                        data_hora_obj = data_hora_val
                    elif isinstance(data_hora_val, str):
                        data_limpa = data_hora_val.strip() 
                        data_hora_obj = datetime.datetime.strptime(data_limpa, "%Y-%m-%d %H:%M:%S")
                    else:
                        continue
                except Exception:
                    continue 

                try:
                    peso_kg_float = float(str(peso_kg_val).replace(',', '.')) 
                except (ValueError, TypeError):
                    continue 

                dados_combinados.append({
                    'Data_Hora_Completa': data_hora_obj,
                    'Peso_KG': peso_kg_float,
                    'Nome': nome_jogador
                })
        return dados_combinados
    except Exception as e:
        import traceback
        traceback.print_exc()
        return f"ERRO FATAL (Openpyxl). Ocorreu um erro. Tipo: {type(e).__name__}"

# =========================================================================
# FERRAMENTA 2: Escrita no Word (LÓGICA ATUALIZADA - 7 DIAS DE REGISTO)
# =========================================================================

def preencher_relatorio_word(dados_limpos: List[Dict[str, Any]], caminho_template: str) -> str:
    """
    Preenche uma tabela Word. Encontra os 7 DIAS DE REGISTO MAIS RECENTES,
    escreve essas datas no cabeçalho da tabela, e preenche os pesos.
    """
    try:
        document = Document(caminho_template)
        
        if not document.tables:
            return "ERRO: O template Word não contém tabelas."
        table = document.tables[0]
        
        jogadores_map = {}
        for i, row in enumerate(table.rows):
            if i == 0: continue 
            nome_jogador_tabela = row.cells[0].text.strip()
            if nome_jogador_tabela:
                jogadores_map[nome_jogador_tabela] = i 

        if not dados_limpos:
            return "ERRO: Não há dados limpos para processar."

        # --- NOVA LÓGICA: Encontrar os 7 dias de registo mais recentes ---
        
        # 1. Obter todas as datas únicas (sem horas) dos registos
        datas_com_registos = set()
        for reg in dados_limpos:
            datas_com_registos.add(reg['Data_Hora_Completa'].date())
            
        # 2. Ordenar as datas únicas (da mais recente para a mais antiga)
        datas_ordenadas = sorted(list(datas_com_registos), reverse=True)
        
        # 3. Selecionar as 7 datas mais recentes (ou menos, se não houver 7)
        ultimas_7_datas_registo = datas_ordenadas[:7]
        
        # 4. Inverter a lista para o relatório (da mais antiga para a mais recente)
        lista_datas_relatorio = sorted(ultimas_7_datas_registo)
        
        print(f"INFO: A gerar relatório para os 7 dias de registo mais recentes.")
        print(f"Datas: {[d.strftime('%d/%m') for d in lista_datas_relatorio]}")

        # --- PREENCHER O CABEÇALHO DO WORD ---
        dias_map = {} 
        header_cells = table.rows[0].cells
        
        for idx_coluna, data_obj in enumerate(lista_datas_relatorio, start=1):
            data_str_coluna = data_obj.strftime('%d/%m') 
            header_cells[idx_coluna].text = data_str_coluna
            dias_map[data_obj] = idx_coluna # Mapeia o objeto data ao índice da coluna

        # --- LÓGICA DE PROCESSAMENTO DE DADOS ---
        pesos_finais = {} 
        for registo in dados_limpos:
            data_hora_registo = registo['Data_Hora_Completa'] 
            data_registo = data_hora_registo.date()
            
            # Apenas processa se a data estiver na nossa lista das 7 datas
            if data_registo not in lista_datas_relatorio:
                continue
            
            chave = (registo['Nome'], data_registo)
            
            if chave not in pesos_finais:
                pesos_finais[chave] = registo
            elif data_hora_registo > pesos_finais[chave]['Data_Hora_Completa']:
                pesos_finais[chave] = registo

        # --- LÓGICA DE PREENCHIMENTO DA TABELA ---
        for (nome, data_dia), registo in pesos_finais.items():
            if nome in jogadores_map and data_dia in dias_map:
                idx_linha = jogadores_map[nome]
                idx_col = dias_map[data_dia] 
                celula = table.rows[idx_linha].cells[idx_col]
                celula.text = "" 
                run = celula.paragraphs[0].add_run(
                    f"{registo['Peso_KG']:.2f}".replace('.', ',')
                )
                run.font.size = Pt(10) 
                run.font.name = 'Calibri' 

        caminho_output = 'relatorio_preenchido.docx'
        document.save(caminho_output)
        return f"Sucesso! Relatório salvo em: {caminho_output}"

    except Exception as e:
        import traceback
        traceback.print_exc()
        return f"ERRO FATAL (Word). Ocorreu um erro. Tipo: {type(e).__name__}: {e}"

# --- FERRAMENTA 1 ---
print("A executar Ferramenta 1: Leitura do Excel...")
ficheiro_excel_original = '/kaggle/input/data-body-weight/dados_peso.xlsx' 
dados_limpos = ler_dados_peso_excel_openpyxl(ficheiro_excel_original)

if not isinstance(dados_limpos, list) or len(dados_limpos) == 0:
    print("Erro na Ferramenta 1. A parar.")
else:
    print(f"Sucesso na Ferramenta 1: {len(dados_limpos)} registos lidos.")

    # --- FERRAMENTA 2 ---
    print("\nA executar Ferramenta 2: Escrita no Word...")

    # <<<--- A MUDANÇA ESTÁ AQUI ---<<<
    ficheiro_template = '/kaggle/input/data-body-weight/template_limpo.docx' 

    # Chama a função de escrita
    resultado_word = preencher_relatorio_word(dados_limpos, ficheiro_template)

    print(f"\n--- Resultado da Ferramenta 2 ---")
    print(resultado_word)

