!pip install ipywidgets
!pip install pymupdf
!pip install python-docx


import os
os.environ["backend"] = "jax"

import tensorflow as tf

import jax.numpy as jnp
from jax import device_put

import keras_cv
import keras
from keras import ops
from keras import backend as K
from keras import layers, models, Model
from keras.callbacks import EarlyStopping, ModelCheckpoint
from keras import mixed_precision

from sklearn.ensemble import RandomForestClassifier
from sklearn.model_selection import train_test_split, StratifiedKFold
from sklearn.metrics import classification_report, confusion_matrix, accuracy_score

from scipy.signal import butter, filtfilt, welch, spectrogram
from scipy.stats import skew, kurtosis
from scipy.signal import lfilter

from mne.preprocessing import ICA
from mne import create_info, EpochsArray
from mne.io import RawArray
from mne.channels import make_standard_montage

from docx import Document
from docx.shared import Inches
from io import BytesIO

import cv2
import pandas as pd
import numpy as np
import shutil
from glob import glob
from tqdm.notebook import tqdm
from joblib import Parallel, delayed
import pywt
import random
import math

import matplotlib.pyplot as plt
import seaborn as sns


print("TensorFlow:", tf.__version__)
print("Keras:", keras.__version__)
print("KerasCV:", keras_cv.__version__)


class CFG:
    verbose = 1  # Verbosity
    pandas_warning = None #Pandas Chained Assignment Warning Mode, default = 'warn'
    seed = 42  # Random seed
    batch_size = 32
    epochs = 20

pd.options.mode.chained_assignment = CFG.pandas_warning
keras.utils.set_random_seed(CFG.seed)

BASE_PATH = "/kaggle/input/hms-harmful-brain-activity-classification"
EEG_DIR = "/tmp/dataset/hms-hbac/numpy_eegs"
os.makedirs(EEG_DIR, exist_ok=True)


df = pd.read_csv(f'{BASE_PATH}/train.csv')
display(df.info())


# As each file points to only one EEG ID, we will select only 1 sample from metadata.
df = df.groupby("eeg_id").head(1).reset_index(drop=True)
df.head(5)


# As we are working on only seizure / non-seizure classification, we will not use default class labels of the dataset.
df['is_seizure'] = df['expert_consensus'] == 'Seizure'

# Train + Valid + Test (We will use train_test_split to split the data)
df['eeg_path'] = f'{BASE_PATH}/train_eegs/'+df['eeg_id'].astype(str)+'.parquet'
# df['eeg2_path'] = f'{EEG_DIR}/'+df['eeg_id'].astype(str)+'.npy'
df['spec_path'] = f'{BASE_PATH}/train_spectrograms/'+df['spectrogram_id'].astype(str)+'.parquet'

display(df.head(2))
print(f"Total Data: {len(df)}")



# Filter for 'Seizure' class and a subset of 'non-Seizure' classes
seizure_df = df[df['expert_consensus'] == 'Seizure']
non_seizure_df = df[df['expert_consensus'] != 'Seizure']

# Undersample non-Seizure data (e.g., 1:1 ratio)
non_seizure_sampled = non_seizure_df.sample(n=len(seizure_df), random_state=CFG.seed)

# Combine and shuffle the new dataset
df = pd.concat([seizure_df, non_seizure_sampled]).sample(frac=1, random_state=CFG.seed)
df.sort_values(by=["eeg_id", "eeg_label_offset_seconds"], ignore_index=True, inplace=True)

class_counts = df['expert_consensus'].value_counts()
display(class_counts)
class_counts.plot(kind='bar', title='Class Distribution')


random_samples = np.random.randint(0, len(df) - 1, size=(10))
for i in random_samples:
    reading_spec = pd.read_parquet(f"{BASE_PATH}/train_spectrograms/{df['spectrogram_id'].iloc[i]}.parquet")
    reading_eeg = pd.read_parquet(f"{BASE_PATH}/train_eegs/{df['eeg_id'].iloc[i]}.parquet")
    print(f"Rows in EEG File: {reading_eeg.shape[0]}\t Rows in Spectrogram File: {reading_spec.shape[0]}")


# Information regarding a single spectrogram
train_df_samples = len(df) # There are total 106800 training samples
print(f"Total Samples: {train_df_samples}")

random_index = np.random.randint(0, train_df_samples - 1)

random_spec_id = df['spectrogram_id'].iloc[random_index]
reading_spec = pd.read_parquet(f"{BASE_PATH}/train_spectrograms/{random_spec_id}.parquet")
random_eeg_id = df['eeg_id'].iloc[random_index]
reading_eeg = pd.read_parquet(f"{BASE_PATH}/train_eegs/{random_eeg_id}.parquet")

print("\nSPECTROGRAM INFORMATION:")
print(f"Shape of spectrogram: {reading_spec.shape}")
print(f"Spectrogram data min: {reading_spec.iloc[:, 1:].values.min()}")
print(f"Spectrogram data max: {reading_spec.iloc[:, 1:].values.max()}")
print(f"Spectrogram data mean: {reading_spec.iloc[:, 1:].values.mean()}")
print("Features of Spectrogram:\n")
display(reading_spec.columns)

print("\nEEG INFORMATION:")
print(f"Shape of EEG: {reading_eeg.shape}")
print(f"EEG data min: {reading_eeg.iloc[:, :-1].values.min()}")
print(f"EEG data max: {reading_eeg.iloc[:, :-1].values.max()}")
print(f"EEG data mean: {reading_eeg.iloc[:, :-1].values.mean()}")
print("Features of Spectrogram:\n")
display(reading_eeg.columns)


import fitz  # PyMuPDF
from PIL import Image
from io import BytesIO

# Folder containing the PDF files
folder_path = f"{BASE_PATH}/example_figures"

# List all files in the folder
pdf_files = [f for f in os.listdir(folder_path) if f.endswith('.pdf')]

# Loop through all PDF files and display them
for pdf_file in pdf_files:
    # Construct the full path of the PDF file
    pdf_path = os.path.join(folder_path, pdf_file)
    
    # Open the PDF file
    doc = fitz.open(pdf_path)
    
    # Loop through the pages and display each one
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)  # Load the page
        pix = page.get_pixmap()  # Convert page to a pixmap (image)

        # Convert to a PIL Image
        img = Image.open(BytesIO(pix.tobytes("png")))

        # Display the image in the notebook
        display(img)



reading_spec = pd.read_parquet(f"{BASE_PATH}/train_spectrograms/{df['spectrogram_id'].iloc[np.random.randint(0, train_df_samples - 1)]}.parquet")

# Step 1: Verify time consistency
def check_time_consistency(spectrogram_data, metadata, spectrogram_id):
    """
    Verify if the 'time' column in the spectrogram data matches the metadata offsets.
    """
    sub_metadata = metadata[metadata["spectrogram_id"] == spectrogram_id]
    metadata_times = sub_metadata["spectrogram_label_offset_seconds"].values
    data_times = spectrogram_data["time"].values

    # Check for missing or mismatched times
    missing_times = set(metadata_times) - set(data_times)
    extra_times = set(data_times) - set(metadata_times)

    print("Time Consistency Check:")
    if missing_times:
        print(f"  Missing time points in spectrogram data: {missing_times}")
    if extra_times:
        print(f"  Extra time points in spectrogram data: {extra_times}")
    if not missing_times and not extra_times:
        print("  All time points match!")
    return missing_times, extra_times

# missing_times, extra_times = check_time_consistency(reading_spec, train_df, random_id)

# Step 2: Check for signal quality across time
def check_signal_quality(df, min_variance=1e-5):
    """
    Analyze the signal quality over time for anomalies or low variance.
    """
    print("\nSignal Quality Check:")
    time_series = df["time"].diff().dropna()

    # Check for irregular time gaps
    if time_series.std() > 0.1:  # Adjust threshold as needed
        print(f"  Irregular time gaps detected! Std Dev: {time_series.std()}")
    else:
        print("  Time gaps are consistent.")

    # Check variance in frequency bands over time
    variances = df.iloc[:, 1:].var(axis=0)  # Exclude 'time' column
    low_variance_cols = variances[variances < min_variance]
    if not low_variance_cols.empty:
        print(f"  Low variance detected in columns: {low_variance_cols.index.tolist()}")
    else:
        print("  All frequency bands have sufficient variance.")

check_signal_quality(reading_spec)
    
import matplotlib.colors as colors

# Step 3: Visualize time-based trends
def plot_spectrogram_over_time(df):
    """
    Plot the spectrogram over time to visually inspect for noise or artifacts.
    """

    # Extract time and spectrogram data
    time = df["time"].values
    spectrogram_data = df.iloc[:, 1:].values  # Excluding the time column

    # Check the min, max, and mean values to understand the scale
    spec_min = spectrogram_data.min()
    spec_max = spectrogram_data.max()
    print(f"Spectrogram data min: {spec_min}")
    print(f"Spectrogram data max: {spec_max}")
    print(f"Spectrogram data mean: {spectrogram_data.mean()}")

    colormaps = ['hsv', "viridis", 'gist_ncar', 'gist_ncar_r', 'gist_rainbow', 'gist_rainbow_r']
    plt.figure(figsize=(12, len(colormaps) * 6))
    for i, cm in enumerate(colormaps):
        plt.subplot(len(colormaps), 1, i + 1)
        plt.imshow(spectrogram_data.T, aspect="auto", cmap=cm,
                   norm=colors.LogNorm(),
                   extent=[time.min(), time.max(), 0, spectrogram_data.shape[0] - 1])
        plt.colorbar(label="Amplitude")
        plt.xlabel("Time (seconds)")
        plt.ylabel("Frequency Bands")
        plt.title(f"Spectrogram Over Time (Color Map: {cm})")
    plt.show()

plot_spectrogram_over_time(reading_spec)


# Bandpass filter function (to remove noise outside 0.5-50 Hz range)
def bandpass_filter(data, lowcut=0.5, highcut=50, fs=200, order=5):
    nyquist = 0.5 * fs
    low = lowcut / nyquist
    high = highcut / nyquist
    b, a = butter(order, [low, high], btype="band")
    return filtfilt(b, a, data, axis=0)


# Load a sample EEG file (update with actual file paths)
reading_eeg = pd.read_parquet(f"{BASE_PATH}/train_eegs/{df['eeg_id'].iloc[np.random.randint(0, train_df_samples - 1)]}.parquet")
eeg_data = reading_eeg.drop(columns=["EKG"])

# Sampling frequency (Hz)
FS = 200

# Analysis Results Dictionary
results = {}

### 1. Check Signal Variance
variances = eeg_data.var()
mean_variance = variances.mean()
std_variance = variances.std()
high_variance_threshold = mean_variance + 2 * std_variance
low_variance_threshold = mean_variance - 2 * std_variance

# Identify electrodes with high/low variance
high_variance_electrodes = variances[variances > high_variance_threshold]
low_variance_electrodes = variances[variances < low_variance_threshold]

results["high_variance"] = high_variance_electrodes
results["low_variance"] = low_variance_electrodes

### 2. Frequency Analysis (FFT)
for column in eeg_data.columns:
    signal = eeg_data[column]
    freq = np.fft.rfftfreq(len(signal), d=1 / FS)  # Frequency axis
    fft_magnitude = np.abs(np.fft.rfft(signal))  # Magnitudes

    # Calculate power in specific frequency bands
    power = { 
        "delta": np.sum(fft_magnitude[(freq >= 0.5) & (freq < 4)]),
        "theta": np.sum(fft_magnitude[(freq >= 4) & (freq < 8)]),
        "alpha": np.sum(fft_magnitude[(freq >= 8) & (freq < 13)]),
        "beta": np.sum(fft_magnitude[(freq >= 13) & (freq < 30)]),
        "gamma": np.sum(fft_magnitude[(freq >= 30) & (freq < 50)]),
        "high_freq_noise": np.sum(fft_magnitude[freq >= 50]),
    }

    # Print power in each band for the first few channels
    if column in eeg_data.columns[:3]:
        print(f"Power for {column}: {power}")
    
    # Record high-frequency noise if it dominates
    if power["high_freq_noise"] > 0.2 * sum(power.values()):
        results[f"{column}_high_freq_noise"] = True
    else:
        results[f"{column}_high_freq_noise"] = False

### 3. Signal Consistency Check
# Compute correlation between electrodes
corr_matrix = eeg_data.corr()
correlated_channels = (corr_matrix > 0.9).sum(axis=1) - 1  # Exclude self-correlation

# Mark channels with high correlation to others (possible redundancy)
results["highly_correlated_electrodes"] = correlated_channels[correlated_channels > 5]

# Visualize correlation matrix
plt.figure(figsize=(30, 30))
sns.set(font_scale=1.5)
sns.heatmap(corr_matrix, annot=True, cmap="coolwarm")
plt.title("Correlation Matrix of EEG Signals")
plt.show()

### 4. Filtered Signal Inspection
for column in eeg_data.columns[:3]:  # Check a few electrodes
    filtered_signal = bandpass_filter(eeg_data[column], 0.5, 50, FS)
    plt.figure(figsize=(10, 3))
    plt.plot(filtered_signal, label=f"Filtered {column}")
    plt.title(f"Filtered EEG Signal: {column}")
    plt.xlabel("Time (samples)")
    plt.ylabel("Amplitude")
    plt.legend()
    plt.show()

### Summary Report
print("\n=== EEG Signal Noise Analysis Summary ===")
print(f"Mean Variance: {mean_variance}")
print(f"High Variance Threshold: {high_variance_threshold}")
print(f"Low Variance Threshold: {low_variance_threshold}")
print(f"Electrodes with High Variance:\n{high_variance_electrodes}")
print(f"Electrodes with Low Variance:\n{low_variance_electrodes}")
print(f"Electrodes with High Variance:\n{results['high_variance']}")
print(f"Electrodes with Low Variance:\n{results['low_variance']}")
print(f"Electrodes with High-Frequency Noise:")
for key, val in results.items():
    if "high_freq_noise" in key and val:
        print(f"  {key.replace('_high_freq_noise', '')}")

print("\nHighly Correlated Electrodes (likely redundant):")
print(results["highly_correlated_electrodes"])



def plot_temporal_signals(eeg_df, num_samples=3, electrodes=["Fp1", "F3", "C3"]):
    """
    Plots temporal EEG signals for a random selection of seizure and non-seizure samples.
    
    Parameters:
        eeg_df (DataFrame): Dataframe containing EEG data and metadata.
        num_samples (int): Number of samples to visualize for seizure and non-seizure each.
        electrodes (list): List of electrode names to plot.
    """
    seizure_samples = eeg_df[eeg_df['is_seizure']].sample(n=num_samples, random_state=42)
    non_seizure_samples = eeg_df[~eeg_df['is_seizure']].sample(n=num_samples, random_state=42)

    # Total number of samples to plot
    total_samples = len(seizure_samples) + len(non_seizure_samples)
    
    # Adjust the number of rows in the subplot grid
    fig, axes = plt.subplots(total_samples, len(electrodes), figsize=(15, 5 * total_samples))

    for idx, sample in enumerate([*seizure_samples.iterrows(), *non_seizure_samples.iterrows()]):
        sample_data = sample[1]
        eeg_path = sample_data['eeg_path']
        eeg_data = pd.read_parquet(eeg_path)

        for col_idx, electrode in enumerate(electrodes):
            signal = eeg_data[electrode]

            # Handle single subplot case when there is only 1 row
            ax = axes[idx, col_idx] if total_samples > 1 else axes[col_idx]
            ax.plot(signal)
            ax.set_title(f"{'Seizure' if sample_data['is_seizure'] else 'Non-Seizure'} - {electrode}")
            ax.set_xlabel("Time (ms)")
            ax.set_ylabel("Amplitude")

    plt.tight_layout()
    plt.savefig("/kaggle/working/Temporal Consistency Analysis.png")
    plt.show()

# Execute the plotting
plot_temporal_signals(df, num_samples=3, electrodes=["Fp1", "F3", "C3"])


def preprocess_signal(signal, clip_threshold=1e6):
    """
    Preprocess the EEG signal by clipping and replacing NaN values.

    Parameters:
        signal (array-like): The raw EEG signal.
        clip_threshold (float): The threshold to clip spikes in the signal.

    Returns:
        array-like: The preprocessed EEG signal.
    """
    # Replace NaN values with the mean of the signal
    signal = np.nan_to_num(signal, nan=np.nanmean(signal))

    # Clip large spikes in the signal
    signal = np.clip(signal, -clip_threshold, clip_threshold)

    return signal

# Update the frequency analysis function
def plot_frequency_analysis_with_preprocessing(eeg_df, num_samples=3, electrodes=["Fp1", "F3", "C3"], fs=200):
    seizure_samples = eeg_df[eeg_df['is_seizure']].sample(n=num_samples, random_state=42)
    non_seizure_samples = eeg_df[~eeg_df['is_seizure']].sample(n=num_samples, random_state=42)

    total_samples = len(seizure_samples) + len(non_seizure_samples)
    fig, axes = plt.subplots(total_samples, len(electrodes), figsize=(15, 5 * total_samples))

    for idx, sample in enumerate([*seizure_samples.iterrows(), *non_seizure_samples.iterrows()]):
        sample_data = sample[1]
        eeg_path = sample_data['eeg_path']
        eeg_data = pd.read_parquet(eeg_path)

        for col_idx, electrode in enumerate(electrodes):
            signal = preprocess_signal(eeg_data[electrode])

            freqs, psd = welch(signal, fs=fs, nperseg=fs * 2)

            ax = axes[idx, col_idx] if total_samples > 1 else axes[col_idx]
            ax.plot(freqs, 10 * np.log10(psd + 1e-12))  # Add small constant to avoid log10(0)
            ax.set_title(f"{'Seizure' if sample_data['is_seizure'] else 'Non-Seizure'} - {electrode}")
            ax.set_xlabel("Frequency (Hz)")
            ax.set_ylabel("Power (dB/Hz)")

    plt.tight_layout()
    plt.savefig("/kaggle/working/Frequency Analysis.png")
    plt.show()

# Run the updated analysis
plot_frequency_analysis_with_preprocessing(df, num_samples=3, electrodes=["Fp1", "F3", "C3"], fs=200)


# Create a folder for saving output if it doesn't exist
output_dir = '/kaggle/working/EEG Analysis'
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

# Create a Word document
doc = Document()
doc.add_heading('EEG Data Analysis', 0)

# Function to add figures to the Word document
def add_figure_to_doc(fig, caption):
    img_stream = BytesIO()
    fig.savefig(img_stream, format='png')
    img_stream.seek(0)  # Rewind to the beginning of the image
    doc.add_picture(img_stream, width=Inches(5))  # Add image to the Word doc
    doc.add_paragraph(caption)


sampling_rate = 200  # Replace with your actual EEG sampling rate

# Channels
channels = eeg_data.columns

# 1. Visual Inspection: Plot raw EEG signals
def plot_raw_signals(eeg_data, channels, duration=5, fs=200):
    time = np.arange(eeg_data.shape[0]) / fs
    fig, axes = plt.subplots(5, 1, figsize=(15, 10))
    for i, channel in enumerate(channels[:5]):  # Plot first 5 channels
        axes[i].plot(time[:fs * duration], eeg_data[channel].iloc[:fs * duration])
        axes[i].set_title(f"Raw EEG Signal - {channel}")
        axes[i].set_xlabel("Time (s)")
        axes[i].set_ylabel("Amplitude")
    plt.tight_layout()
    add_figure_to_doc(fig, 'Raw EEG Signal for First 5 Channels')
    plt.show()

plot_raw_signals(eeg_data, channels)

# 2. Power Spectral Density (PSD) Analysis
def plot_psd(eeg_data, channels, fs=200):
    fig, axes = plt.subplots(5, 1, figsize=(15, 10))
    for i, channel in enumerate(channels[:5]):  # Plot first 5 channels
        freqs, psd = welch(eeg_data[channel], fs=fs, nperseg=fs * 2)
        axes[i].semilogy(freqs, psd)
        axes[i].set_title(f"Power Spectral Density - {channel}")
        axes[i].set_xlabel("Frequency (Hz)")
        axes[i].set_ylabel("Power (dB/Hz)")
    plt.tight_layout()
    add_figure_to_doc(fig, 'Power Spectral Density for First 5 Channels')
    plt.show()

plot_psd(eeg_data, channels)

# 3. Check for Missing Data or Flat Channels
def check_missing_flat_channels(eeg_data):
    channel_stats = eeg_data.describe().T
    flat_channels = channel_stats.loc[channel_stats['std'] == 0].index.tolist()
    doc.add_paragraph(f"Flat Channels: {flat_channels}")
    print(f"Flat Channels: {flat_channels}")
    return flat_channels

flat_channels = check_missing_flat_channels(eeg_data)

# 4. Band-Specific Analysis
def compute_band_power(eeg_data, fs=200):
    bands = {"Delta": (0.5, 4), "Theta": (4, 8), "Alpha": (8, 13), "Beta": (13, 30), "Gamma": (30, 40)}
    band_power = {}
    for band, (low, high) in bands.items():
        band_power[band] = eeg_data.apply(lambda x: welch(x, fs=fs, nperseg=fs * 2)[1][
            (low <= welch(x, fs=fs, nperseg=fs * 2)[0]) & 
            (welch(x, fs=fs, nperseg=fs * 2)[0] <= high)].sum(), axis=0)
    band_power = pd.DataFrame(band_power)
    band_power.plot(kind='bar', figsize=(12, 6), title="Relative Band Power")
    plt.ylabel("Power")
    plt.xlabel("Channels")
    plt.tight_layout()
    add_figure_to_doc(plt, 'Band-Specific Power Analysis')
    plt.show()

compute_band_power(eeg_data)

# 5. Independent Component Analysis (ICA)
def run_ica(eeg_data, sampling_rate=200, highpass_freq=1.0):
    n_channels = len(channels)
    info = create_info(ch_names=list(channels), sfreq=sampling_rate, ch_types=["eeg"] * n_channels)
    
    # Create random 3D coordinates for each channel
    # This is just for visualization purposes, ideally you'd use a proper montage
    info.set_montage('standard_1020')  # Standard system of EEG recordings
    
    # Create a RawArray from the EEG data
    raw_array = RawArray(eeg_data.values.T, info)
    
    # Apply high-pass filter to the data
    raw_array.filter(l_freq=highpass_freq, h_freq=None)  # High-pass filter with 1.0 Hz lower bound
    
    ica = ICA(n_components=n_channels, random_state=42)
    ica.fit(raw_array)  # Now fitting to RawArray
    
    # Plot ICA components and save the figure manually
    ica_fig = ica.plot_components(picks=None, show=False)  # This returns a figure
    ica_fig.savefig(os.path.join(output_dir, 'ica_components.png'))  # Save ICA components plot to file
    
    # Add the ICA components plot to the document
    add_figure_to_doc(ica_fig, 'Independent Component Analysis (ICA) Components')
    return ica

ica = run_ica(eeg_data)

# 6. Correlation Between Channels
def plot_channel_correlation(eeg_data):
    corr_matrix = eeg_data.corr()
    fig, ax = plt.subplots(figsize=(10, 8))
    cax = ax.imshow(corr_matrix, cmap='viridis', interpolation='none')
    fig.colorbar(cax, ax=ax, label='Correlation')
    ax.set_title("Channel Correlation Matrix")
    ax.set_xlabel("Channels")
    ax.set_ylabel("Channels")
    plt.tight_layout()
    add_figure_to_doc(fig, 'Channel Correlation Matrix')
    plt.show()

plot_channel_correlation(eeg_data)

# 7. Signal-to-Noise Ratio (SNR)
def compute_snr(eeg_data):
    signal_power = (eeg_data ** 2).mean(axis=0)
    noise_power = eeg_data.diff().dropna() ** 2
    noise_power = noise_power.mean(axis=0)
    snr = 10 * np.log10(signal_power / noise_power)
    doc.add_paragraph("Signal-to-Noise Ratio (SNR) for each channel:")
    doc.add_paragraph(snr.to_string())
    print("SNR for each channel:")
    print(snr)
    return snr

snr = compute_snr(eeg_data)

# 8. Epoch-Level Noise Analysis
def check_epoch_noise(eeg_data, epoch_duration=2, fs=200):
    epoch_samples = fs * epoch_duration
    num_epochs = eeg_data.shape[0] // epoch_samples
    variances = []
    for i in range(num_epochs):
        epoch = eeg_data.iloc[i * epoch_samples:(i + 1) * epoch_samples]
        variances.append(epoch.var(axis=0))
    variances = pd.DataFrame(variances)
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.boxplot(variances.values)
    ax.set_title("Epoch-Level Variance Across Channels")
    ax.set_xlabel("Channels")
    ax.set_ylabel("Variance")
    plt.tight_layout()
    add_figure_to_doc(fig, 'Epoch-Level Variance Across Channels')
    plt.show()

check_epoch_noise(eeg_data)

# 9. Statistical Properties
def plot_statistical_properties(eeg_data):
    stats = {
        "Mean": eeg_data.mean(),
        "Std": eeg_data.std(),
        "Skewness": eeg_data.apply(skew),
        "Kurtosis": eeg_data.apply(kurtosis),
    }
    stats_df = pd.DataFrame(stats)
    stats_df.plot(kind='bar', figsize=(15, 6), subplots=True, layout=(2, 2), title="Statistical Properties")
    plt.tight_layout()
    add_figure_to_doc(plt, 'Statistical Properties')
    plt.show()

plot_statistical_properties(eeg_data)

# Save the Word document
doc.save(os.path.join(output_dir, 'EEG_Analysis_Report.docx'))



# Function to compute SNR for one file
def compute_snr_for_file(eeg_file, clip_threshold=1e6):
    eeg_data = pd.read_parquet(eeg_file)
    eeg_data.drop(columns=["EKG"], inplace=True)
    
    # Handle NaNs or Infs
    if eeg_data.isnull().values.any() or np.isinf(eeg_data.values).any():
        eeg_data = eeg_data.fillna(0)
        eeg_data = eeg_data.replace([np.inf, -np.inf], 0)
    
    eeg_data = eeg_data.clip(-clip_threshold, clip_threshold)  # Clip extreme values

    # Compute signal and noise power
    signal_power = (eeg_data ** 2).mean(axis=0)
    if np.isnan(signal_power).any():
        print(f"Warning: NaN values detected in signal power for {eeg_file}.")
    
    noise_power = eeg_data.diff().dropna() ** 2
    noise_power = noise_power.mean(axis=0)
    if np.isnan(noise_power).any():
        print(f"Warning: NaN values detected in noise power for {eeg_file}.")
    
    # Add small epsilon to avoid log(0)
    epsilon = 1e-10
    snr = 10 * np.log10((signal_power + epsilon) / (noise_power + epsilon))
    
    return snr

# Randomly sample EEG files for analysis
sample_files = df.sample(100, random_state=CFG.seed)['eeg_path']

# Compute SNR for each file
snr_results = []
for file in sample_files:
    snr = compute_snr_for_file(file)
    snr_results.append(snr)

# Aggregate SNR results across files
snr_df = pd.DataFrame(snr_results).mean(axis=0)
print("Average SNR for Channels:")
print(snr_df)

# Filter channels with high SNR
snr_threshold = 0  # Adjust based on results
high_snr_channels = snr_df[snr_df > snr_threshold].index.tolist()
print("High-SNR Channels:", high_snr_channels)


# Select a balanced subset of seizure and non-seizure data
seizure_files = df[df['is_seizure']]['eeg_path'].sample(50, random_state=CFG.seed)
non_seizure_files = df[~df['is_seizure']]['eeg_path'].sample(50, random_state=CFG.seed)
files_to_process = pd.concat([seizure_files, non_seizure_files])

# Prepare data for classification
X, y = [], []
for file in files_to_process:
    eeg_data = pd.read_parquet(file)[high_snr_channels]
    X.append(eeg_data.mean(axis=0))  # Aggregate features (e.g., mean)
    y.append(df.loc[df['eeg_path'] == file, 'is_seizure'].values[0])

# Convert to DataFrame
X = pd.DataFrame(X, columns=high_snr_channels)
y = np.array(y)

# Train-test split
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=CFG.seed)

# Train a Random Forest Classifier
rf_model = RandomForestClassifier(random_state=CFG.seed)
rf_model.fit(X_train, y_train)

# Feature importance
importances = pd.Series(rf_model.feature_importances_, index=high_snr_channels)
importances.sort_values(ascending=False).plot(kind='bar', figsize=(10, 6), title="Channel Importance")
plt.savefig("Feature Importance.png")
plt.show()


# Selected top channels based on analysis
SELECTED_CHANNELS = ['Fp1', 'O2', 'T6', 'Fz', 'F4', 'T3', 'Cz', 'T5', 'C4', 'P3']

# Define thresholds
LOW_VARIANCE_THRESHOLD = 0.001
HIGH_VARIANCE_THRESHOLD = 10

def wavelet_denoise(data, wavelet="db4", level=2):
    coeffs = pywt.wavedec(data, wavelet, axis=0)
    threshold = np.median(np.abs(coeffs[-1])) / 0.6745  # Universal threshold
    coeffs = [pywt.threshold(c, threshold, mode="soft") for c in coeffs]
    return pywt.waverec(coeffs, wavelet, axis=0)

def calculate_snr(signal):
    mean_signal = np.mean(signal)
    noise = signal - mean_signal
    snr = 10 * np.log10(np.var(signal) / np.var(noise))
    return snr

# Define a function to process a single eeg_id
def evaluate_eeg(eeg_id):
    eeg_path = f"{BASE_PATH}/train_eegs/{eeg_id}.parquet"
    try:
        # Load EEG data
        eeg = pd.read_parquet(eeg_path)
        
        # Select relevant channels
        chan_eeg = eeg[SELECTED_CHANNELS]
        
        # Replace NaN and infinite values
        chan_eeg = chan_eeg.replace([np.inf, -np.inf], 0).fillna(0)

        # Apply bandpass filtering
        chan_eeg = bandpass_filter(chan_eeg)

        # Wavelet denoising
        chan_eeg = wavelet_denoise(chan_eeg)

        # Normalize data with safeguard for zero standard deviation
        stds = chan_eeg.std(axis=0)
        stds[stds == 0] = 1e-8  # Replace zero stds with a small constant
        normalized_eeg = (chan_eeg - chan_eeg.mean(axis=0)) / stds
        normalized_eeg = np.nan_to_num(normalized_eeg)  # Ensure no NaN or inf values

        # Variance checks
        variances = normalized_eeg.var(axis=0)
        if (variances < LOW_VARIANCE_THRESHOLD).any() or (variances > HIGH_VARIANCE_THRESHOLD).any():
            # print(f"Discarded {eeg_id} due to variance thresholds after preprocessing.")
            return None

        # SNR checks
        snr_values = np.apply_along_axis(calculate_snr, axis=0, arr=normalized_eeg)
        if (snr_values < 0).any():
            # print(f"Discarded {eeg_id} due to negative SNR after preprocessing.")
            return None
        
        # Undersample to 6,000 rows
        # The EEGs were recored at 200 samples / second
        # 200 * 30 = 6000 Samples
        # Therefore we will pick first 6000 samples which will consolidate 30 seconds of recording
        sub_eeg = normalized_eeg[:6000, :]
        
        # Convert to NumPy and save
        # eeg_np = sub_eeg.to_numpy(dtype='float', na_value=0)
        return sub_eeg  # Return eeg_id if successfully processed
    
    except Exception as e:
        # print(f"Error processing {eeg_id}: {e}")
        return None

def process_and_save_eeg(eeg_id):
    processed = evaluate_eeg(eeg_id)
    if processed is not None:
        np.save(f"{EEG_DIR}/{eeg_id}.npy", processed)
        return eeg_id  # Return ID if successfully processed
    return None

if __name__ == "__main__":
    # Collect all processed EEG IDs
    results = Parallel(n_jobs=-1, backend="loky")(
        delayed(process_and_save_eeg)(eeg_id)
        for eeg_id in tqdm(df["eeg_id"], desc="Processing EEG files")
    )

    # Filter out None values and save successfully processed IDs
    accepted_eegs = [eeg_id for eeg_id in results if eeg_id is not None]
    print(f"Successfully processed {len(accepted_eegs)} EEG files.")



accepted_df = df[df["eeg_id"].isin(accepted_eegs)]
class_counts = accepted_df['expert_consensus'].value_counts()
patients = accepted_df["patient_id"].value_counts()
display(class_counts)


# Encode the "is_seizure" label (True -> 1, False -> 0)
accepted_df["is_seizure"] = accepted_df["is_seizure"].astype(int)

# Convert to NumPy arrays
eeg_ids = accepted_df["eeg_id"].to_numpy()
labels = accepted_df["is_seizure"].to_numpy()

# Number of folds
n_splits = 5

skf = StratifiedKFold(n_splits=n_splits, shuffle=True, random_state=CFG.seed)
splits = list(skf.split(eeg_ids, labels))

# Placeholder to store datasets for each fold
fold_datasets = []

# Function to load EEG data from .npy files
def load_eeg(eeg_id, label):
    eeg_id_str = tf.strings.as_string(eeg_id)
    eeg_path = tf.strings.join([EEG_DIR, tf.strings.join([eeg_id_str, ".npy"], separator="")], separator="/")
    eeg_data = tf.numpy_function(np.load, [eeg_path], tf.float64)
    eeg_data = tf.cast(eeg_data, tf.float32)
    eeg_data = tf.ensure_shape(eeg_data, [6000, 10])
    return eeg_data, label

# Function to apply data augmentation
def augment_data(eeg_data, label):
    # Add Gaussian noise
    noise = tf.random.normal(tf.shape(eeg_data), mean=0.0, stddev=0.02)
    eeg_data = eeg_data + noise

    # Random time shifting
    time_shift = tf.random.uniform([], minval=-50, maxval=50, dtype=tf.int32)
    eeg_data = tf.roll(eeg_data, shift=time_shift, axis=0)

    return eeg_data, label

# Function to preprocess the data
def preprocess_data(eeg_data, label):
    eeg_data = tf.expand_dims(eeg_data, axis=-1)  # Add a channel dimension: (6000, 10) -> (6000, 10, 1)
    return eeg_data, label

# Function to create TensorFlow datasets for a fold
def create_fold_dataset(train_idx, test_idx):
    # Get train and test EEG IDs and labels
    train_ids, train_labels = eeg_ids[train_idx], labels[train_idx]
    test_ids, test_labels = eeg_ids[test_idx], labels[test_idx]

    # Create TensorFlow datasets
    train_dataset = tf.data.Dataset.from_tensor_slices((train_ids, train_labels))
    test_dataset = tf.data.Dataset.from_tensor_slices((test_ids, test_labels))

    # Map the loading and preprocessing functions
    train_dataset = train_dataset.map(
        lambda eeg_id, label: load_eeg(eeg_id, label), num_parallel_calls=tf.data.AUTOTUNE
    )
    train_dataset = train_dataset.map(preprocess_data, num_parallel_calls=tf.data.AUTOTUNE)
    test_dataset = test_dataset.map(
        lambda eeg_id, label: load_eeg(eeg_id, label), num_parallel_calls=tf.data.AUTOTUNE
    )
    test_dataset = test_dataset.map(preprocess_data, num_parallel_calls=tf.data.AUTOTUNE)

    # Shuffle, batch, and prefetch
    batch_size = CFG.batch_size
    train_dataset = train_dataset.shuffle(buffer_size=1000).batch(batch_size).prefetch(tf.data.AUTOTUNE)
    test_dataset = test_dataset.batch(batch_size).prefetch(tf.data.AUTOTUNE)

    return train_dataset, test_dataset

# Loop through each fold and create datasets
for i, (train_idx, test_idx) in enumerate(splits):
    print(f"Creating datasets for Fold {i + 1}")
    train_dataset, test_dataset = create_fold_dataset(train_idx, test_idx)
    fold_datasets.append((train_dataset, test_dataset))

print(f"Successfully created datasets for {n_splits} folds.")






def create_cnn_lstm(input_shape=(6000, 10, 1), num_classes=2):
    model = models.Sequential(name="CNN-BiLSTM-Model")

    # Input Layer
    model.add(layers.Input(shape=input_shape, name="Input-Layer"))
    
    # 1. CNN Block (Spatial Feature Extraction)
    model.add(layers.Conv2D(name="Convolution-1", filters=32, kernel_size=(1, 5), activation='relu', padding='same'))  # Added padding
    model.add(layers.BatchNormalization(name="Normalization-1"))
    model.add(layers.MaxPooling2D(name="Pooling-1", pool_size=(1, 2)))  # Reduces the width by a factor of 2

    model.add(layers.Conv2D(name="Convolution-2", filters=64, kernel_size=(1, 3), activation='relu', padding='same'))  # Added padding
    model.add(layers.BatchNormalization(name="Normalization-2"))
    model.add(layers.MaxPooling2D(name="Pooling-2", pool_size=(1, 2)))  # Reduces the width by a factor of 2
    
    # After Conv2D and MaxPooling layers, the shape is reduced to:
    # Height = 6000 (unchanged), Width = 10 -> (10 / 2 / 2) = 2, Depth = 64
    # So, the shape after this block is (6000, 2, 64)
    
    # 2. Temporal Downsampling (Reshape for LSTM)
    # The output shape after CNN and Pooling is (6000, 2, 64), so we reshape it to (6000, 128)
    model.add(layers.Reshape(name="Reshape-Conv-Output", target_shape=(6000, 128)))  # 6000 time steps, 128 features
    
    # 3. LSTM Block (Temporal Dependency Capture)
    model.add(layers.Bidirectional(layers.LSTM(128, return_sequences=True), name="Bi-LSTM-1"))
    model.add(layers.Dropout(0.2, name="Dropout-1"))
    model.add(layers.Bidirectional(layers.LSTM(64), name="Bi-LSTM-2"))

    # 4. Fully Connected Layers (Classification)
    model.add(layers.Dense(64, name="Dense-1", activation='relu'))
    model.add(layers.Dropout(0.2, name="Dropout-2"))
    model.add(layers.Dense(32, name="Dense-2", activation='relu'))
    model.add(layers.Dense(num_classes, name="Output-Layer", activation='softmax'))

    # Compile the model
    model.compile(optimizer='adam', loss='sparse_categorical_crossentropy', metrics=['accuracy'])
    return model

# Instantiate the model
model = create_cnn_lstm()
model.summary()



ckpt_cb = tf.keras.callbacks.ModelCheckpoint(
    'cnn_lstm.keras',                # File path to save the model
    monitor="val_loss",
    save_best_only=True,             # Save regardless of validation loss
    save_weights_only=False,         # Save the full model (not just weights)
    save_freq='epoch',               # Save at the end of each epoch (which is the default)
    verbose=0                        # Display a message when saving
)

# Early stopping callback
erst_cb = EarlyStopping(
    monitor="val_loss",  # Monitor validation loss
    patience=3,          # Number of epochs to wait before stopping
    restore_best_weights=True,  # Restore the model with the best weights
    verbose=CFG.verbose
)


# Hyperparameters
epochs = 10

# Loop through each fold
for fold_idx, (train_dataset, test_dataset) in enumerate(fold_datasets):
    print(f"\nTraining Fold {fold_idx + 1}/{len(fold_datasets)}")
    
    # Split train_dataset into train and validation datasets
    validation_split = 0.2  # 20% for validation
    total_train_size = len(train_dataset)
    val_size = int(total_train_size * validation_split)
    
    train_dataset = train_dataset.skip(val_size)
    val_dataset = train_dataset.take(val_size)
    
    # Train the model
    training = model.fit(
        train_dataset,
        validation_data=val_dataset,
        epochs=epochs,
        callbacks=[ckpt_cb, erst_cb],
        verbose=CFG.verbose,
    )

    # Extract loss and accuracy values from the history object
    train_loss = training.history['loss']
    val_loss = training.history['val_loss']
    train_acc = training.history['accuracy']
    val_acc = training.history['val_accuracy']
    
    # Plotting the training and validation loss
    plt.figure(figsize=(12, 6))
    plt.subplot(1, 2, 1)
    plt.plot(range(1, len(train_loss) + 1), train_loss, label='Training Loss')
    plt.plot(range(1, len(val_loss) + 1), val_loss, label='Validation Loss')
    plt.title('Loss over Epochs')
    plt.xlabel('Epochs')
    plt.ylabel('Loss')
    plt.legend()
    
    # Plotting the training and validation accuracy
    plt.subplot(1, 2, 2)
    plt.plot(range(1, len(train_acc) + 1), train_acc, label='Training Accuracy')
    plt.plot(range(1, len(val_acc) + 1), val_acc, label='Validation Accuracy')
    plt.title('Accuracy over Epochs')
    plt.xlabel('Epochs')
    plt.ylabel('Accuracy')
    plt.legend()
    
    # Display the plots
    plt.tight_layout()
    plt.show()
    
    # Evaluate on the testing dataset
    test_loss, test_accuracy = model.evaluate(test_dataset, verbose=0)
    print(f"Fold {fold_idx + 1} - Test Accuracy: {test_accuracy:.4f}")


# Collect true and predicted labels for the test set across all folds
y_true_all = []
y_pred_all = []

print("Evaluation Running...")

# Loop through each fold's test dataset
for fold_idx, (train_dataset, test_dataset) in enumerate(fold_datasets):
    print(f"\nEvaluating Fold {fold_idx + 1}/{len(fold_datasets)}")
    
    y_true = []
    y_pred = []
    
    for eeg_data, labels in test_dataset:
        predictions = model.predict(eeg_data, verbose=0)
        y_true.extend(labels.numpy())
        y_pred.extend(tf.argmax(predictions, axis=1).numpy())
    
    # Append results for all folds
    y_true_all.extend(y_true)
    y_pred_all.extend(y_pred)

print("Evaluation Ends!")



cm = confusion_matrix(y_true_all, y_pred_all)
# Plot confusion matrix heatmap
plt.figure(figsize=(8, 5))
sns.heatmap(cm, annot=True, fmt='d', cmap='Blues')
plt.xlabel("Predicted")
plt.ylabel("True")
plt.show()

# Print the average accuracy & specificity model acheived
TN, FP, FN, TP = cm.ravel()
precision = TP / (TP + FP)
sensitivity = TP / (TP + FN)
specificity = TN / (TN + FP)
acc = accuracy_score(y_true_all, y_pred_all)
print(f"\033[1mModel Precision: {round(precision, 2) * 100}%\033[0m")
print(f"\033[1mModel Sensitivity: {round(sensitivity, 2) * 100}%\033[0m")
print(f"\033[1mModel Specificity: {round(specificity, 2) * 100}%\033[0m")
print(f"\033[1mModel Mean Accuracy: {round(acc, 2) * 100}%\n\033[0m")

# Generate a classification report
print("\t\t\tCLASSIFICATION REPORT:")
print(classification_report(y_true_all, y_pred_all, target_names=["Non-Seizure", "Seizure"]))



model.save("cnn_lstm.keras")

