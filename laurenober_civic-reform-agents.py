# Core ADK imports
from google.adk.agents import Agent, LlmAgent, SequentialAgent, ParallelAgent, LoopAgent
from google.adk.models.google_llm import Gemini
from google.adk.runners import Runner
from google.adk.sessions import DatabaseSessionService
from google.adk.memory import InMemoryMemoryService
from google.adk.tools import FunctionTool, google_search, AgentTool
from google.genai import types

# Standard library
import os
import json
import re
import logging
from datetime import datetime
from typing import Dict, List, Any, Optional
import glob

# Install required packages if needed
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')  # Non-interactive backend
except ImportError:
    import subprocess
    import sys
    print("ğŸ“¦ Installing required packages...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", 
                          "python-docx", "matplotlib", "--quiet"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')
    print("âœ… Packages installed")

print("âœ… All imports successful!")



def cleanup_output_folder(working_dir: str = "/kaggle/working"):
    """
    Clean up the output folder by removing old policy briefs.
    Keeps the folder tidy and prevents accumulation of old files.
    """
    patterns_to_clean = [
        "policy_brief_*.docx",
        "policy_analysis_evaluation.json"
    ]
    
    files_removed = 0
    for pattern in patterns_to_clean:
        for filepath in glob.glob(f"{working_dir}/{pattern}"):
            try:
                os.remove(filepath)
                files_removed += 1
            except Exception as e:
                print(f"âš ï¸�  Could not remove {filepath}: {e}")
    
    if files_removed > 0:
        print(f"ğŸ§¹ Cleaned up {files_removed} old files from output folder")
    
    return files_removed

# Run cleanup at start
cleanup_output_folder()


from kaggle_secrets import UserSecretsClient

try:
    GOOGLE_API_KEY = UserSecretsClient().get_secret("GOOGLE_API_KEY")
    os.environ["GOOGLE_API_KEY"] = GOOGLE_API_KEY
    print("âœ… Gemini API key configured successfully.")
except Exception as e:
    print(f"ğŸ”’ Authentication Error: {e}")
    print("Please add 'GOOGLE_API_KEY' to your Kaggle secrets.")


# Configure retry options for API calls
retry_config = types.HttpRetryOptions(
    attempts=5,
    exp_base=7,
    initial_delay=1,
    http_status_codes=[429, 500, 503, 504],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# Global configuration
APP_NAME = "civic_policy_analyzer"
USER_ID = "policy_researcher"
MODEL_NAME = "gemini-2.5-flash-lite"

print("âœ… Configuration complete!")
print(f"   - Model: {MODEL_NAME}")
print(f"   - App: {APP_NAME}")
print(f"   - Retry attempts: {retry_config.attempts}")


def verify_claim_with_source(claim: str, source_url: str = "") -> Dict[str, Any]:
    """
    Verify a statistical claim and assess its credibility.
    
    Args:
        claim: The claim to verify (e.g., "17% increase in voter turnout")
        source_url: URL of the source making the claim
        
    Returns:
        Dictionary with verification results
    """
    result = {
        "claim": claim,
        "source": source_url,
        "credibility_score": 0.0,
        "verification_status": "unverified",
        "concerns": []
    }
    
    # Extract domain
    if source_url:
        domain = source_url.split('/')[2] if '/' in source_url else source_url
        domain = domain.lower()
        
        # High credibility sources
        high_cred = ['.gov', '.edu', 'scholar.google', 'jstor', 'pubmed', 
                     'nature.com', 'science.org', 'census.gov', 'pewresearch.org']
        # Medium credibility
        med_cred = ['nytimes.com', 'washingtonpost.com', 'reuters.com', 'apnews.com', 'bbc.com']
        # Advocacy sources (lower credibility for unbiased data)
        advocacy = ['fairvote', 'represent.us', 'commoncause']
        
        if any(trusted in domain for trusted in high_cred):
            result["credibility_score"] = 0.9
            result["verification_status"] = "high_confidence"
        elif any(trusted in domain for trusted in med_cred):
            result["credibility_score"] = 0.7
            result["verification_status"] = "medium_confidence"
        elif any(adv in domain for adv in advocacy):
            result["credibility_score"] = 0.5
            result["verification_status"] = "advocacy_source"
            result["concerns"].append("Source is advocacy organization - may have bias")
        else:
            result["credibility_score"] = 0.4
            result["verification_status"] = "unknown_source"
            result["concerns"].append("Source credibility unknown")
    
    # Check for vague language
    vague_terms = ['approximately', 'around', 'about', 'some studies', 'may increase', 'could lead to']
    if any(term in claim.lower() for term in vague_terms):
        result["concerns"].append("Claim contains vague language")
    
    # Check for specific numbers
    has_specific_number = bool(re.search(r'\d+(?:\.\d+)?%', claim))
    if not has_specific_number:
        result["concerns"].append("Claim lacks specific statistical evidence")
    
    return result


def extract_verified_metrics(text: str, source_verification: bool = True) -> Dict[str, Any]:
    """
    Extract metrics from text with verification flags.
    
    Args:
        text: Research text containing statistics
        source_verification: Whether to flag unverified claims
        
    Returns:
        Dictionary with verified and unverified metrics separated
    """
    metrics = {
        "verified_statistics": [],
        "unverified_claims": [],
        "percentages": [],
        "years": [],
        "sample_sizes": []
    }
    
    # Extract percentages with context
    sentences = text.split('.')
    for sentence in sentences:
        # Look for percentages with citation patterns
        if '%' in sentence:
            has_citation = any(pattern in sentence.lower() for pattern in 
                             ['according to', 'study by', 'found that', 'research from', 'data from'])
            has_hedge = any(hedge in sentence.lower() for hedge in 
                          ['may', 'might', 'could', 'approximately', 'around', 'some studies'])
            
            percentage_match = re.search(r'(\d+(?:\.\d+)?)\s*%', sentence)
            if percentage_match:
                stat_info = {
                    "value": float(percentage_match.group(1)),
                    "context": sentence.strip(),
                    "has_citation": has_citation,
                    "has_hedge_language": has_hedge
                }
                
                if has_citation and not has_hedge:
                    metrics["verified_statistics"].append(stat_info)
                else:
                    metrics["unverified_claims"].append(stat_info)
                
                metrics["percentages"].append(float(percentage_match.group(1)))
    
    # Extract years
    years = re.findall(r'\b(19\d{2}|20\d{2})\b', text)
    metrics["years"] = sorted(list(set(years)))
    
    # Extract sample sizes (n=X patterns)
    samples = re.findall(r'n\s*=\s*(\d+(?:,\d{3})*)', text.lower())
    metrics["sample_sizes"] = [int(s.replace(',', '')) for s in samples]
    
    return metrics


def create_policy_visualization(data: Dict[str, Any], chart_type: str = "comparison") -> str:
    """
    Create a visualization for policy data.
    
    Args:
        data: Data to visualize
        chart_type: Type of chart (comparison, timeline, distribution)
        
    Returns:
        Filepath to saved chart
    """
    plt.figure(figsize=(10, 6))
    
    if chart_type == "comparison" and "pros" in data and "cons" in data:
        # Create a simple comparison chart
        categories = ['Pros', 'Cons']
        values = [len(data.get("pros", [])), len(data.get("cons", []))]
        colors = ['#2ecc71', '#e74c3c']
        
        plt.bar(categories, values, color=colors, alpha=0.7, edgecolor='black')
        plt.title('Policy Arguments Balance', fontsize=14, fontweight='bold')
        plt.ylabel('Number of Arguments', fontsize=12)
        plt.grid(axis='y', alpha=0.3)
        
    elif chart_type == "timeline" and "implementations" in data:
        # Create timeline of implementations
        implementations = data["implementations"]
        years = [impl.get("year", 2000) for impl in implementations]
        successes = [1 if impl.get("success", False) else 0 for impl in implementations]
        
        plt.scatter(years, successes, s=100, alpha=0.6, c=successes, cmap='RdYlGn')
        plt.title('Historical Implementation Timeline', fontsize=14, fontweight='bold')
        plt.xlabel('Year', fontsize=12)
        plt.ylabel('Success (1) / Failure (0)', fontsize=12)
        plt.grid(True, alpha=0.3)
        
    elif chart_type == "credibility" and "sources" in data:
        # Source credibility distribution
        sources = data["sources"]
        cred_levels = ['High\n(.gov, .edu)', 'Medium\n(Major news)', 'Low\n(Unknown)']
        counts = [
            sum(1 for s in sources if s.get("credibility", 0) >= 0.8),
            sum(1 for s in sources if 0.6 <= s.get("credibility", 0) < 0.8),
            sum(1 for s in sources if s.get("credibility", 0) < 0.6)
        ]
        colors = ['#27ae60', '#f39c12', '#e74c3c']
        
        plt.bar(cred_levels, counts, color=colors, alpha=0.7, edgecolor='black')
        plt.title('Source Credibility Distribution', fontsize=14, fontweight='bold')
        plt.ylabel('Number of Sources', fontsize=12)
        plt.grid(axis='y', alpha=0.3)
    
    # Save chart
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filepath = f"/kaggle/working/policy_chart_{chart_type}_{timestamp}.png"
    plt.tight_layout()
    plt.savefig(filepath, dpi=150, bbox_inches='tight')
    plt.close()
    
    return filepath


print("âœ… Enhanced tools created:")
print("   - verify_claim_with_source()")
print("   - extract_verified_metrics()")
print("   - create_policy_visualization()")


# RESEARCH AGENT - More focused, less redundant
research_agent = Agent(
    name="ResearchAgent",
    model=Gemini(model=MODEL_NAME, retry_options=retry_config),
    instruction="""You are a policy research specialist. Gather ONLY the most important, 
    well-sourced facts about the policy topic.
    
    Requirements:
    - Find 3-5 KEY findings (not 10-15)
    - Each finding must cite a specific, credible source
    - Prioritize government data, peer-reviewed studies, major news outlets
    - AVOID advocacy group websites unless corroborated by credible sources
    - Flag any claims that seem questionable or lack verification
    
    Format each finding as:
    **Finding:** [Clear, concise statement]
    **Source:** [Specific source with credibility indicator]
    **Verification:** [High/Medium/Low confidence based on source]
    
    Be concise. Quality over quantity.""",
    tools=[google_search],
    output_key="research_findings"
)

# HISTORICAL PRECEDENT AGENT - More data-driven
historical_agent = Agent(
    name="HistoricalPrecedentAgent",
    model=Gemini(model=MODEL_NAME, retry_options=retry_config),
    instruction="""You identify where this policy has been implemented with MEASURABLE outcomes.
    
    Find 3-4 BEST examples (not 10+) that have:
    - Clear implementation date
    - Measurable outcomes (voter turnout %, cost savings, etc.)
    - Current status (still active? repealed?)
    - Credible source for the data
    
    For each example provide:
    **Location & Year:** [e.g., "Maine, USA - 2018"]
    **Outcome Metrics:** [Specific numbers with sources]
    **Current Status:** [Active/Repealed/Modified]
    **Key Lesson:** [One sentence takeaway]
    
    CRITICAL: Verify claims. If a source says "voter turnout increased 17%" - find the 
    actual study. Flag unverified claims.""",
    tools=[google_search],
    output_key="historical_precedents"
)

# METRICS AGENT - Enhanced with verification
metrics_agent = Agent(
    name="MetricsEvidenceAgent",
    model=Gemini(model=MODEL_NAME, retry_options=retry_config),
    instruction="""You extract and VERIFY quantitative evidence from research.
    
    For each statistic you find, assess:
    1. Is it from a credible source? (.gov, .edu, peer-reviewed)
    2. Is there hedge language? ("may increase", "some studies suggest")
    3. Is the sample size adequate?
    4. Is it corroborated by multiple sources?
    
    Output format:
    ## Verified Statistics (High Confidence)
    - [Stat with source and verification notes]
    
    ## Claimed Statistics (Needs Verification)  
    - [Stat with concerns listed]
    
    ## Data Quality Assessment
    - Overall confidence: High/Medium/Low
    - Main concerns: [List any issues]
    - Recommendations: [What additional verification needed]
    
    Be skeptical of advocacy group statistics unless verified independently.""",
    tools=[
        FunctionTool(verify_claim_with_source),
        FunctionTool(extract_verified_metrics)
    ],
    output_key="metrics_analysis"
)

# SYNTHESIS AGENT - More structured, less redundant
synthesis_agent = Agent(
    name="SynthesisAgent",
    model=Gemini(model=MODEL_NAME, retry_options=retry_config),
    instruction="""You compile research into a concise, balanced policy brief.
    
    CRITICAL RULES:
    - Keep each section CONCISE (2-4 paragraphs max)
    - NO redundancy - don't repeat information across sections
    - Lead with verified facts, flag unverified claims
    - Be balanced but honest about data quality
    - Include ONLY the most impactful arguments on each side
    
    Structure:
    
    ## Executive Summary
    (3-4 sentences: What it is, key finding, recommendation)
    
    ## Policy Overview  
    (2-3 paragraphs: What's proposed, who's affected, core objectives)
    
    ## Evidence Base
    (3-4 key findings with credibility assessment)
    
    ## Pros & Cons
    ### Strengths (Top 3-4 arguments with evidence)
    ### Concerns (Top 3-4 arguments with evidence)
    
    ## Key Stakeholders
    (Who supports/opposes - names and reasoning in 1 paragraph each)
    
    ## Implementation Roadmap
    (3-4 key phases with realistic timelines)
    
    ## Critical Risks
    (Top 3-4 concerns from Devil's Advocate)
    
    ## Feasibility Assessment
    (Political: High/Med/Low | Technical: High/Med/Low | with justification)
    
    ## Recommendations
    (3-5 clear action items)
    
    AVOID: Long lists, repetition, unverified claims presented as fact, 
    excessive detail on minor points.""",
    tools=[],
    output_key="final_policy_brief"
)

# DEVIL'S ADVOCATE - More targeted
devils_advocate_agent = Agent(
    name="DevilsAdvocateAgent",
    model=Gemini(model=MODEL_NAME, retry_options=retry_config),
    instruction="""Challenge the analysis to identify the TOP 3-5 most critical weaknesses.
    
    Focus on:
    - Unverified claims being presented as facts
    - Cherry-picked data or biased sources
    - Ignored downsides or risks  
    - Unrealistic timelines or cost estimates
    - Logical inconsistencies
    
    For each concern:
    **Issue:** [Clear description]
    **Impact:** [Why it matters]
    **Severity:** [Critical/Moderate/Minor]
    **Mitigation:** [How to address it]
    
    Be constructive. Identify the MOST important issues, not every tiny flaw.
    
    If the analysis is genuinely strong, say so and note what makes it credible.""",
    tools=[],
    output_key="critical_analysis"
)

print("âœ… Agents created with better prompts")


# Advocacy and resistance agents (keeping these simpler)
advocacy_agent = Agent(
    name="AdvocacyResearchAgent",
    model=Gemini(model=MODEL_NAME, retry_options=retry_config),
    instruction="""Identify the TOP 3-5 key advocates and their main arguments.
    
    Focus on:
    - Prominent politicians/officials supporting it
    - Major organizations backing it
    - Their 3-4 strongest arguments with evidence
    
    Be concise. Quality over quantity.""",
    tools=[google_search],
    output_key="advocacy_research"
)

resistance_agent = Agent(
    name="ResistanceAnalysisAgent",
    model=Gemini(model=MODEL_NAME, retry_options=retry_config),
    instruction="""Identify the TOP 3-5 key opponents and main obstacles.
    
    Focus on:
    - Prominent politicians/groups opposing it
    - Their 3-4 strongest objections with evidence
    - Political and practical implementation challenges
    
    Assess political feasibility: High/Medium/Low with clear reasoning.""",
    tools=[google_search],
    output_key="resistance_analysis"
)

implementation_agent = Agent(
    name="ImplementationPathwayAgent",
    model=Gemini(model=MODEL_NAME, retry_options=retry_config),
    instruction="""Create a realistic 3-4 phase implementation roadmap.
    
    For each phase:
    - Clear objective (1 sentence)
    - Timeline estimate (be realistic)
    - 3-4 key actions
    - Success criteria
    
    Be practical. Don't create 20-step plans - focus on critical path.""",
    tools=[],
    output_key="implementation_roadmap"
)

# Build orchestrator
parallel_research = ParallelAgent(
    name="ParallelResearchTeam",
    sub_agents=[research_agent, historical_agent, advocacy_agent, resistance_agent]
)

sequential_analysis = SequentialAgent(
    name="SequentialAnalysisTeam",
    sub_agents=[metrics_agent, implementation_agent]
)

synthesis_loop = LoopAgent(
    name="SynthesisQualityLoop",
    sub_agents=[synthesis_agent, devils_advocate_agent],
    max_iterations=2
)

orchestrator_agent = SequentialAgent(
    name="PolicyAnalysisOrchestrator",
    sub_agents=[parallel_research, sequential_analysis, synthesis_loop]
)

print("âœ… Streamlined orchestrator created")


db_url = "sqlite:///civic_policy_analysis.db"
session_service = DatabaseSessionService(db_url=db_url)
memory_service = InMemoryMemoryService()

runner = Runner(
    agent=orchestrator_agent,
    app_name=APP_NAME,
    session_service=session_service,
    memory_service=memory_service
)

print("âœ… Session & Memory configured")


# ğŸ�¯ SET YOUR POLICY QUESTION HERE:
policy_question = "Should we implement universal healthcare?"

# ğŸ”„ OPTIONAL: Resume a previous session (leave as None for new session)
resume_session_id = None  # Example: "policy_analysis_20241119_213045"

print(f"\nğŸš€ Starting policy analysis for: {policy_question}")
print("="*70)

# Use existing session or create new one
if resume_session_id:
    session_id = resume_session_id
    print(f"ğŸ“‚ Resuming session: {session_id}")
else:
    session_id = f"policy_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    print(f"ğŸ†• Starting new session: {session_id}")

# Run the analysis
print("â�³ Agents working...\n")

# run_debug returns a list of events
responses = await runner.run_debug(
    policy_question,
    session_id=session_id
)

print("\nâœ… Analysis complete!\n")

# Debug: Show what we received
print("="*70)
print("ğŸ”� AGENT RESPONSE DEBUG")
print("="*70)
print(f"Total responses received: {len(responses)}")

# Extract the final policy brief - look for the SYNTHESIS agent output specifically
policy_brief = None

# Search through responses for the synthesis agent output
for response in reversed(responses):  # Start from end
    response_text = None
    
    # Extract text from response
    if hasattr(response, 'text') and response.text:
        response_text = response.text
    elif hasattr(response, 'content'):
        if hasattr(response.content, 'parts') and response.content.parts:
            response_text = response.content.parts[0].text
        elif hasattr(response.content, 'text'):
            response_text = response.content.text
    
    # Check if this looks like the synthesis output (has multiple sections)
    if response_text and len(response_text) > 1000:
        # Look for policy brief structure markers
        if any(marker in response_text for marker in ['## Executive Summary', '## Policy Overview', '## Evidence']):
            policy_brief = response_text
            print(f"âœ… Found synthesis output: {len(policy_brief)} characters")
            break

# If we didn't find a structured brief, try to reconstruct from agent outputs
if not policy_brief or len(policy_brief) < 1000:
    print("âš ï¸�  Synthesis output not found in standard format. Reconstructing from agent outputs...")
    
    agent_outputs = []
    for response in responses:
        if hasattr(response, 'text') and response.text and len(response.text) > 100:
            agent_outputs.append(response.text)
    
    # Combine the substantial agent outputs
    if agent_outputs:
        policy_brief = "\n\n".join(agent_outputs[-5:])  # Last 5 substantial outputs
        print(f"âœ… Reconstructed from {len(agent_outputs)} agent outputs")

if not policy_brief:
    policy_brief = "Error: Could not extract policy brief. The agent may not have completed successfully."
    print(f"â�Œ Warning: {policy_brief}")
else:
    print(f"\nğŸ“„ Policy Brief Length: {len(policy_brief)} characters")
    print(f"ğŸ“Š Policy Brief Preview (first 500 chars):")
    print("-"*70)
    print(f"\nğŸ’¾ Session saved as: {session_id}")
    print(f"   To resume this analysis later, use: resume_session_id = '{session_id}'")
    print(policy_brief[:500] + "..." if len(policy_brief) > 500 else policy_brief)
    print("-"*70)


def create_summary_charts(policy_brief: str, session_id: str) -> Dict[str, Any]:
    """
    Create visualizations based on the policy brief content.
    Returns dictionary of chart types to image objects (BytesIO).
    """
    from io import BytesIO
    
    charts = {}
    brief_lower = policy_brief.lower()
    
    # Chart 1: Pros vs Cons Count
    try:
        # Count pros and cons by looking for bullet points in those sections
        pros_section = re.search(r'###?\s*(?:pros|strengths|benefits).*?(?=###|\n##|\Z)', 
                                policy_brief, re.IGNORECASE | re.DOTALL)
        cons_section = re.search(r'###?\s*(?:cons|concerns|challenges|risks).*?(?=###|\n##|\Z)', 
                                policy_brief, re.IGNORECASE | re.DOTALL)
        
        pros_count = len(re.findall(r'^\s*[-*â€¢]', pros_section.group(0), re.MULTILINE)) if pros_section else 0
        cons_count = len(re.findall(r'^\s*[-*â€¢]', cons_section.group(0), re.MULTILINE)) if cons_section else 0
        
        if pros_count > 0 or cons_count > 0:
            plt.figure(figsize=(8, 5))
            categories = ['Strengths', 'Concerns']
            values = [pros_count, cons_count]
            colors = ['#27ae60', '#e74c3c']
            
            bars = plt.bar(categories, values, color=colors, alpha=0.7, edgecolor='black', linewidth=2)
            plt.title('Policy Analysis Balance', fontsize=16, fontweight='bold', pad=20)
            plt.ylabel('Number of Arguments', fontsize=12)
            plt.ylim(0, max(values) * 1.2 if max(values) > 0 else 10)
            
            # Add value labels on bars
            for bar in bars:
                height = bar.get_height()
                plt.text(bar.get_x() + bar.get_width()/2., height,
                        f'{int(height)}',
                        ha='center', va='bottom', fontsize=14, fontweight='bold')
            
            plt.grid(axis='y', alpha=0.3, linestyle='--')
            plt.tight_layout()
            
            # Save to BytesIO instead of file
            img_buffer = BytesIO()
            plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
            img_buffer.seek(0)
            plt.close()
            charts['balance'] = img_buffer
    except Exception as e:
        print(f"âš ï¸�  Could not create balance chart: {e}")
    
    # Chart 2: Data Quality Indicators
    try:
        has_verified = 'verified' in brief_lower or 'high confidence' in brief_lower
        has_gov_sources = '.gov' in brief_lower or '.edu' in brief_lower
        has_statistics = bool(re.search(r'\d+%', policy_brief))
        has_years = bool(re.search(r'\b(19|20)\d{2}\b', policy_brief))
        
        indicators = ['Verified\nClaims', 'Gov/Edu\nSources', 'Statistical\nData', 'Historical\nData']
        values = [
            1 if has_verified else 0,
            1 if has_gov_sources else 0,
            1 if has_statistics else 0,
            1 if has_years else 0
        ]
        
        if sum(values) > 0:
            plt.figure(figsize=(8, 5))
            colors_quality = ['#27ae60' if v else '#95a5a6' for v in values]
            
            bars = plt.bar(indicators, values, color=colors_quality, alpha=0.7, edgecolor='black', linewidth=2)
            plt.title('Evidence Quality Indicators', fontsize=16, fontweight='bold', pad=20)
            plt.ylabel('Present (1) / Absent (0)', fontsize=12)
            plt.ylim(0, 1.3)
            
            # Add checkmarks or X marks
            for i, (bar, val) in enumerate(zip(bars, values)):
                symbol = 'âœ“' if val else 'âœ—'
                color = '#27ae60' if val else '#e74c3c'
                plt.text(bar.get_x() + bar.get_width()/2., 0.5,
                        symbol, ha='center', va='center', 
                        fontsize=28, color=color, fontweight='bold')
            
            plt.grid(axis='y', alpha=0.3, linestyle='--')
            plt.tight_layout()
            
            # Save to BytesIO instead of file
            img_buffer = BytesIO()
            plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
            img_buffer.seek(0)
            plt.close()
            charts['quality'] = img_buffer
    except Exception as e:
        print(f"âš ï¸�  Could not create quality chart: {e}")
    
    # Chart 3: Feasibility Assessment
    try:
        feasibility_section = re.search(r'feasibility.*?(?=\n##|\Z)', 
                                       policy_brief, re.IGNORECASE | re.DOTALL)
        if feasibility_section:
            text = feasibility_section.group(0).lower()
            
            # Determine feasibility levels
            political = 'high' if 'political' in text and 'high' in text else \
                       'low' if 'political' in text and 'low' in text else 'medium'
            technical = 'high' if 'technical' in text and 'high' in text else \
                       'low' if 'technical' in text and 'low' in text else 'medium'
            
            categories = ['Political\nFeasibility', 'Technical\nFeasibility']
            values = [
                3 if political == 'high' else 2 if political == 'medium' else 1,
                3 if technical == 'high' else 2 if technical == 'medium' else 1
            ]
            colors_feas = ['#27ae60' if v == 3 else '#f39c12' if v == 2 else '#e74c3c' for v in values]
            
            plt.figure(figsize=(8, 5))
            bars = plt.bar(categories, values, color=colors_feas, alpha=0.7, edgecolor='black', linewidth=2)
            plt.title('Implementation Feasibility', fontsize=16, fontweight='bold', pad=20)
            plt.ylabel('Feasibility Level', fontsize=12)
            plt.ylim(0, 3.5)
            plt.yticks([1, 2, 3], ['Low', 'Medium', 'High'])
            
            # Add labels
            labels = [political.capitalize(), technical.capitalize()]
            for bar, label in zip(bars, labels):
                height = bar.get_height()
                plt.text(bar.get_x() + bar.get_width()/2., height,
                        label, ha='center', va='bottom', 
                        fontsize=12, fontweight='bold')
            
            plt.grid(axis='y', alpha=0.3, linestyle='--')
            plt.tight_layout()
            
            # Save to BytesIO instead of file
            img_buffer = BytesIO()
            plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
            img_buffer.seek(0)
            plt.close()
            charts['feasibility'] = img_buffer
    except Exception as e:
        print(f"âš ï¸�  Could not create feasibility chart: {e}")
    
    return charts


def export_enhanced_policy_brief(
    policy_question: str,
    policy_brief: str,
    session_id: str,
    add_visualizations: bool = True
) -> str:
    """
    Export policy brief to beautifully formatted DOCX with visualizations.
    """
    # Create visualizations first
    charts = {}
    if add_visualizations:
        print("ğŸ“Š Creating visualizations...")
        charts = create_summary_charts(policy_brief, session_id)
        if charts:
            print(f"   âœ… Created {len(charts)} charts")
    
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Title Page
    title = doc.add_heading('POLICY ANALYSIS BRIEF', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Policy Question
    question_para = doc.add_paragraph()
    question_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    question_run = question_para.add_run(policy_question)
    question_run.font.size = Pt(16)
    question_run.font.bold = True
    question_run.font.color.rgb = RGBColor(0, 102, 204)
    
    # Metadata
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(
        f"\nGenerated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n"
        f"Multi-Agent AI Policy Analysis System\n"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # Spacing
    
    # Add key visualizations at the top if available
    if charts:
        viz_heading = doc.add_heading('Key Analysis Metrics', level=1)
        for run in viz_heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Add charts directly from BytesIO objects
        for chart_name, chart_buffer in charts.items():
            try:
                doc.add_picture(chart_buffer, width=Inches(5.5))
                # Center the image
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()  # Spacing
            except Exception as e:
                print(f"âš ï¸�  Could not add chart {chart_name}: {e}")
        
        doc.add_page_break()
    
    # Parse and format policy brief content
    current_section = None
    lines = policy_brief.split('\n')
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue
        
        # Major section headers (##)
        if line_stripped.startswith('## '):
            header_text = line_stripped.replace('## ', '').strip()
            heading = doc.add_heading(header_text, level=1)
            for run in heading.runs:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 51, 102)
            current_section = header_text.lower()
            
        # Subsection headers (###)
        elif line_stripped.startswith('### '):
            subheader_text = line_stripped.replace('### ', '').strip()
            subheading = doc.add_heading(subheader_text, level=2)
            for run in subheading.runs:
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(51, 51, 51)
                
        # Bullet points
        elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
            bullet_text = line_stripped[2:].strip()
            bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
            bullet_para.paragraph_format.left_indent = Inches(0.25)
            bullet_para.paragraph_format.space_before = Pt(3)
            bullet_para.paragraph_format.space_after = Pt(3)
            
        # Numbered lists
        elif re.match(r'^\d+\.', line_stripped):
            list_text = re.sub(r'^\d+\.\s*', '', line_stripped)
            list_para = doc.add_paragraph(list_text, style='List Number')
            list_para.paragraph_format.left_indent = Inches(0.25)
            
        # Bold emphasis (**text**)
        elif '**' in line_stripped:
            para = doc.add_paragraph()
            parts = line_stripped.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Odd indices are between **
                    run = para.add_run(part)
                    run.bold = True
                else:
                    para.add_run(part)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            
        # Regular paragraphs
        else:
            para = doc.add_paragraph(line_stripped)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.15
    
    # Add footer with methodology
    doc.add_page_break()
    footer_heading = doc.add_heading('Methodology & Data Quality', level=1)
    for run in footer_heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    footer_section = doc.add_paragraph()
    footer_run = footer_section.add_run(
        "This policy analysis was generated using a coordinated multi-agent AI system with 8 specialized agents:\n\n"
        "â€¢ Research Agent - Factual information gathering\n"
        "â€¢ Historical Precedent Agent - Past implementation analysis\n"
        "â€¢ Advocacy Analysis Agent - Supporter identification\n"
        "â€¢ Resistance Analysis Agent - Opposition mapping\n"
        "â€¢ Metrics Verification Agent - Data validation\n"
        "â€¢ Implementation Planning Agent - Roadmap creation\n"
        "â€¢ Critical Review Agent - Risk assessment\n"
        "â€¢ Synthesis Agent - Final brief compilation\n\n"
        "All statistical claims have been assessed for credibility. Sources from government agencies "
        "(.gov, .edu) and peer-reviewed research receive higher confidence ratings than advocacy "
        "organization claims.\n\n"
        "âš ï¸�  Important: Please verify all facts and sources before using for official purposes.\n\n"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(60, 60, 60)
    
    # Session info
    session_para = doc.add_paragraph()
    session_run = session_para.add_run(f"Session ID: {session_id}")
    session_run.font.size = Pt(8)
    session_run.font.color.rgb = RGBColor(150, 150, 150)
    session_run.italic = True
    
    # Save document
    safe_filename = re.sub(r'[^\w\s-]', '', policy_question.lower())
    safe_filename = re.sub(r'[-\s]+', '-', safe_filename)
    filename = f"policy_brief_{safe_filename}.docx"
    filepath = f"/kaggle/working/{filename}"
    
    doc.save(filepath)
    return filepath

# Export the brief
print("\nğŸ“� Exporting policy brief to DOCX...")
output_file = export_enhanced_policy_brief(
    policy_question=policy_question,
    policy_brief=policy_brief,
    session_id=session_id,
    add_visualizations=True
)

print(f"\nâœ… Policy brief exported successfully")
print(f"ğŸ“„ File: {output_file}")
print(f"\nğŸ’¡ This brief includes:")
print("   â€¢ Executive summary and key findings")
print("   â€¢ Visual charts showing analysis balance")
print("   â€¢ Evidence quality indicators")
print("   â€¢ Verified statistics with source assessment")
print("   â€¢ Balanced pros and cons analysis")
print("   â€¢ Implementation roadmap with feasibility")
print("   â€¢ Critical risk assessment")
print("   â€¢ Professional formatting for presentations")
print(f"\nğŸ�™ï¸�  Ready for: Podcasts, policy meetings, advocacy campaigns")


def evaluate_brief_quality(policy_brief: str) -> Dict[str, Any]:
    """
    Evaluate the quality of the generated policy brief.
    """
    brief_lower = policy_brief.lower()
    
    results = {
        "balance_score": 0.0,
        "completeness_score": 0.0,
        "data_quality_score": 0.0,
        "conciseness_score": 0.0,
        "issues": [],
        "strengths": []
    }
    
    # Check balance
    has_pros = any(word in brief_lower for word in 
                   ['pros:', 'strengths:', 'benefits:', 'advantages:'])
    has_cons = any(word in brief_lower for word in 
                   ['cons:', 'concerns:', 'challenges:', 'risks:'])
    has_advocates = any(word in brief_lower for word in 
                       ['advocates', 'supporters', 'proponents'])
    has_opponents = any(word in brief_lower for word in 
                       ['opponents', 'critics', 'opposition'])
    
    balance_elements = [has_pros, has_cons, has_advocates, has_opponents]
    results["balance_score"] = sum(balance_elements) / len(balance_elements)
    
    # Check completeness
    required = ['executive summary', 'overview', 'evidence', 'implementation', 
                'recommendation', 'feasibility']
    sections_found = sum(1 for section in required if section in brief_lower)
    results["completeness_score"] = sections_found / len(required)
    
    # Check data quality indicators
    has_verified = 'verified' in brief_lower or 'credibility' in brief_lower
    has_sources = any(source in brief_lower for source in ['.gov', '.edu', 'study', 'research'])
    has_numbers = bool(re.search(r'\d+%', policy_brief))
    
    data_quality_elements = [has_verified, has_sources, has_numbers]
    results["data_quality_score"] = sum(data_quality_elements) / len(data_quality_elements)
    
    # Check conciseness (character count per section - rough heuristic)
    word_count = len(policy_brief.split())
    if word_count < 1000:
        results["issues"].append("Brief may be too short")
    elif word_count > 2500:
        results["issues"].append("Brief may be too long - aim for 1500-2500 words")
    else:
        results["strengths"].append("Good length (1000-2500 words)")
    
    results["conciseness_score"] = min(1.0, 2000 / max(word_count, 1))
    
    # Overall quality
    results["overall_quality"] = (
        results["balance_score"] * 0.35 +
        results["completeness_score"] * 0.25 +
        results["data_quality_score"] * 0.25 +
        results["conciseness_score"] * 0.15
    )
    
    # Add interpretations
    if results["balance_score"] >= 0.75:
        results["strengths"].append("Well-balanced analysis")
    else:
        results["issues"].append("Lacks balance in presenting multiple viewpoints")
    
    if results["data_quality_score"] >= 0.7:
        results["strengths"].append("Good data verification practices")
    else:
        results["issues"].append("Needs better source verification")
    
    return results

# Run evaluation
eval_results = evaluate_brief_quality(policy_brief)

print("\n" + "="*70)
print("ğŸ“Š QUALITY ASSESSMENT")
print("="*70)
print(f"\nâš–ï¸�  Balance Score:        {eval_results['balance_score']:.0%}")
print(f"âœ… Completeness Score:   {eval_results['completeness_score']:.0%}")
print(f"ğŸ“Š Data Quality Score:   {eval_results['data_quality_score']:.0%}")
print(f"ğŸ“� Conciseness Score:    {eval_results['conciseness_score']:.0%}")
print(f"\nğŸ�¯ Overall Quality:      {eval_results['overall_quality']:.0%}")

if eval_results['strengths']:
    print("\nâœ¨ Strengths:")
    for strength in eval_results['strengths']:
        print(f"   â€¢ {strength}")

if eval_results['issues']:
    print("\nâš ï¸�  Areas for Improvement:")
    for issue in eval_results['issues']:
        print(f"   â€¢ {issue}")

print("\n" + "="*70)
print("âœ… Analysis complete- Check your output file.")
print("="*70)

